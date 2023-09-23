from tkinter import *
import tkinter as tk
from tkinter import ttk
import sqlite3 as sq
from datetime import datetime, timedelta
from tkinter.ttk import Combobox
from tkinter.scrolledtext import ScrolledText
from tkinter import scrolledtext, messagebox
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
import os
import re
import random
from docx.enum.section import WD_SECTION, WD_ORIENT


from docxtpl import DocxTemplate
from docxcompose.composer import Composer
import subprocess

import data_base
import vaccinations

data = {
    "patient": {
        'name': '',
        'birth_date': '',
        'gender': '',
        'amb_cart': '',
        'patient_district': '',
        'address': ''},
    "doctor": {
        'doctor_name': '',
        'manager': '',
        'ped_div': '',
        'doctor_district': ''},

    'certificate': {
        'type_certificate': '',
    }

}

all_data = {
    'sport_section': ('баскетболом', 'волейболом', 'вольной борьбой', 'гандболом', 'греблей', 'каратэ',
                      'легкой атлетикой', 'музыкой', 'плаванием в бассейне', 'спортивной гимнастикой', 'танцами',
                      'теннисом', 'тхэквондо', "ушу", 'фигурным катанием', 'фигурным катанием', 'футболом', 'хоккеем',
                      'шахматами', 'шашками', '+ СОРЕВНОВАНИЯ'),
    "health": {
        "group": ("1", "2", "3", "4"),
        "physical": ("Основная", "Подготовительная", "СМГ", "ЛФК"),
        "regime": ("общий", "ортопедический", "зрительный", "щадящий", "охранительный"),
        "diet": ("А", "Б", "Д", "М", 'П', 'ПП', 'Ц'),
        "desk": ("по росту", "1", "2", "3", "средний ряд")
    },
    "diagnosis": (("<< КАРДИОЛОГИЯ >>", "САС:", "ООО", "ДХЛЖ", "НК0", "НБПНПГ", "ФСШ", "ВПС:", "ДМЖП", "ДМПП"),
                  ("<< ЛОГОПЕДИЯ >>", "ОНР", "ФНР", "ФФНР", "ЗРР", "стертая дизартрия",
                   "ур.р.р.", "1", "2", "3"),
                  ("<< ОФТАЛЬМОЛОГИЯ >>", "Спазм аккомодации", "Миопия", "Гиперметропия",
                   "слабой степени", "средней степени", "тяжелой степени", "OD", "OS", "OU", "с ast"),
                  ("<< ОРТОПЕДИЯ >>", "Нарушение осанки", "Сколиотическая осанка", "Плоскостопие", "ПВУС",
                   "ИС:", "левосторонняя", "правосторонняя", "кифотическая", "грудная", "поясничная",
                   "грудо-поясничная", "деформация позвоночника", "ГПС", "1 ст.", "2 ст.", "3 ст."),
                  ("<< ЛОР >>", "ГА", "ГНМ", "хр. тонзиллит", "1ст", "2ст", "3ст"),
                  ("<< ПРОЧЕЕ >>", "рецидивирующие заболевания ВДП", "Атопический дерматит", "Кариес",
                   "угроза", "БРА", "хр. гастрит", "СДН", "ВСД")),

    'place': ('Средняя школа (гимназия)',
              'Детское Дошкольное Учреждение',
              'Колледж (техникум)',
              'На кружки и секции',
              'В стационар',
              'По месту требования'),

    'type': ('По выздоровлении',
             'Годовой медосмотр',
             'На кружки и секции',
             'ЦКРОиР',
             'Оформление в ДДУ / СШ / ВУЗ',
             'Об отсутствии контактов',
             'В детский лагерь',
             'О нуждаемости в сан-кур лечении',
             "Может работать по специальности...",
             'Об усыновлении (удочерении)',
             "Бесплатное питание",
             "Об обслуживании в поликлинике"),

    "all_info": {

        "По выздоровлении": {
            "additional_medical_information": "| Осмотрен на чесотку, педикулез, микроспорию |",
            "recommendation": "Освобождение от занятий физкультурой и плаванием \nв бассейне на семь дней",
            "validity_period": "5 дней"},

        "Годовой медосмотр": {
            "additional_medical_information": "Рост _____ см; Вес _____ кг; Vis OD/OS = __________; АД ________\n"
                                              "| Осмотрен на чесотку, педикулез, микроспорию |",
            "diagnosis": "Группа здоровья: _ ; Группа по физкультуре: _ ;",
            "recommendation": "Режим _; Стол _; ",
            "date_of_issue": "now"},

        "На кружки и секции": {
            "place_of_requirement": "Для занятия ",
            "diagnosis": "Не имеется медицинских противопоказаний, "
                         "включенных в перечень медицинских противопоказаний для занятия ",
            "date_of_issue": "now",
            "validity_period": "1 год"},

        "Бесплатное питание": {
            "place_of_requirement": "Управление социальной защиты Первомайского района",
            "additional_medical_information": "Ребенок находится на искусственном вскармливании. "
                                              "Получает питание по возрасту",
            "diagnosis": "Соматически здоров",
            "recommendation": "Рекомендован рацион питания ребенку в соответствии с возрастом и примерным месячным "
                              "набором продуктов питания согласно приложения №1 к постановлению Министерства труда "
                              "и социальной защиты РБ и МЗРБ от 13.03.2012 № 37/20",
            "date_of_issue": "now",
            "validity_period": "6 месяцев"},

        "ЦКРОиР": {
            "place_of_requirement": "Для логопедической комиссии (ЦКРОиР)",
            "additional_medical_information": "Ребенок от _____ беременности, _____ родов. \n"
                                              "Течение беременности: без особенностей\n"
                                              "При рождении: вес ________ гр, рост _______ см, Апгар _______ \n"
                                              "К году: зубов ______ , рост ______ см, вес _______ гр\n"
                                              "Держать голову в ______ мес, ползать ______ мес, сидеть ______ мес,\n"
                                              " стоять ____ мес,  ходить _____ мес, говорить _____ мес, "
                                              "первый зуб _____ мес\n"
                                              "Невролог – ____________________________________________________ \n"
                                              "Окулист – _____________________________________________________ \n"
                                              "ЛОР – _________________________________________________________ \n"
                                              "Логопед – _____________________________________________________ ",
            "diagnosis": f"{'_' * 55}",
            "date_of_issue": "now",
            "validity_period": "1 год"},

        "Оформление в ДДУ / СШ / ВУЗ": {
            "additional_medical_information": "Рост _____ см; Вес _____ кг; Vis OD/OS = __________; АД ________\n"
                                              "| Осмотрен на чесотку, педикулез, микроспорию |\n"
                                              "Данные о профилактических прививках прилагаются",
            "diagnosis": f"Группа здоровья: _ ; Группа по физкультуре: _ ;",
            "recommendation": "Режим _; Стол _;",
            "date_of_issue": "now",
            "validity_period": "1 год"},

        "Об отсутствии контактов": {
            "additional_medical_information": "В контакте с инфекционными больными в течение 35 дней не был\n"
                                              "| Осмотрен на чесотку, педикулез, микроспорию |",
            "diagnosis": "На момент осмотра соматически здоров",
            "recommendation": "Может находиться в детском коллективе",
            "date_of_issue": "now",
            "validity_period": "5 дней"},

        "О нуждаемости в сан-кур лечении": {
            "place_of_requirement": "О нуждаемости в санаторно-курортном лечении",
            "date_of_issue": "now",
            "validity_period": "1 год"},

        "В детский лагерь": {
            "place_of_requirement": "В лагерь",
            "additional_medical_information": "Рост _____ см; Вес _____ кг; АД ________\n"
                                              "| Осмотрен на чесотку, педикулез, микроспорию |\n"
                                              "В контакте с инфекционными больными в течение 35 дней не был\n"
                                              "Данные о профилактических прививках прилагаются",
            "diagnosis": f"diagnosis\nГруппа здоровья: _ ; Группа по физкультуре: _ ;\n"
                         "На момент осмотра соматически здоров",
            "recommendation": "Режим _; Стол _;",
            "date_of_issue": "now",
            "validity_period": "5 дней"},

        "Может работать по специальности...": {
            "place_of_requirement": "Проведение обязательного предварительного / внеочередного медицинского осмотра",
            "diagnosis": "Годен к работе по специальности: ",
            "date_of_issue": "now",
            "validity_period": "До следующего обязательного периодического медицинского осмотра"},

        "Об обслуживании в поликлинике": {
            "place_of_requirement": "По месту требования",
            "diagnosis": "Ребенок обслуживается в УЗ '19-я городская детская поликлиника с ___________'",
            "date_of_issue": "now",
            "validity_period": "1 год"},

        "Об усыновлении (удочерении)": {
            "place_of_requirement": "По месту требования",
            "past_illnesses": "При рождении: Рост ______ см; Вес ______ кг; Апгар ________ ; \n"
                              "Семейно-генеалогический анамнез: (отягощен / не отягощен / нет данных)\n",
            "additional_medical_information": "Рост _____ см; Вес _____ кг;\n"
                                              "Осмотры специалистов:\n"
                                              f"Педиатр: дата {'_' * 10} д/з: _________________________________\n"
                                              f"хирург: дата {'_' * 10} д/з: __________________________________\n"
                                              f"офтальмолог: дата {'_' * 10} д/з: _____________________________\n"
                                              f"оториноларинголог: дата {'_' * 10} д/з: _______________________\n"
                                              f"стоматолог: дата {'_' * 10} д/з: ______________________________\n"
                                              f"невролог: дата {'_' * 10} д/з: ________________________________\n"
                                              f"психиатр-нарколог: дата {'_' * 10} д/з: _______________________\n"
                                              f"логопед: дата {'_' * 10} д/з: _________________________________\n"
                                              f"Лабораторные анализы:\n"
                                              f"общий анализ крови: дата {'_' * 10} закл-е: __________________\n"
                                              f"общий анализ мочи: дата {'_' * 10} закл-е: ____________________\n"
                                              f"анализ крови на ВИЧ: дата {'_' * 10} закл-е: __________________\n"
                                              f"Hbs-Ag: дата {'_' * 10} закл-е: _______________________________\n"
                                              f"RW: дата {'_' * 10} закл-е: ___________________________________\n"
                                              "Данные о профилактических прививках прилагаются",
            "diagnosis": f"diagnosis\n"
                         f"Физическое развитие (выше- ниже-) среднее, (дис-) гармоничное\n"
                         f"НПР - соответствует возрасту",
            "date_of_issue": "now",
            "validity_period": "1 год"
        }
    }}

render_data = dict()


def append_info(info):
    for key_0 in info:
        if key_0 == 'text_size':
            data[key_0] = info.get(key_0)
        else:
            for key_1 in info.get(key_0):
                data[key_0][key_1] = info[key_0].get(key_1)
    ask_type_certificate()


def ask_type_certificate():
    def select_type_certificate(event):
        print(event.widget)
        num = ''
        for i in str(event.widget).split('.!')[-1]:
            if i.isdigit():
                num += i
        data['certificate']['type_certificate'] = all_data.get('type')[int(num) - 2]
        type_cert_root.destroy()
        type_cert_root.quit()
        editing_certificate()

    type_cert_root = Toplevel()
    type_cert_root.title('Выбор справки')
    type_cert_root.config(bg='white')

    Label(type_cert_root, text='Какую справку создать?\n',
          font=('Comic Sans MS', data.get('text_size')), bg='white').grid()
    for text in all_data.get('type'):
        lbl_0 = Label(type_cert_root, text=text,
                      font=('Comic Sans MS', data.get('text_size')), border=1, compound='left',
                      bg='#f0fffe', relief='ridge')
        lbl_0.grid(ipadx=2, ipady=2, padx=2, pady=2, sticky='ew')
        lbl_0.bind('<Button-1>', select_type_certificate)

    type_cert_root.mainloop()


def editing_certificate():
    destroy_elements = dict()
    edit_cert_root = Toplevel()

    edit_cert_root.title(f'Редактирование справки')
    edit_cert_root.config(bg='white')
    type_certificate = data['certificate'].get('type_certificate')

    Label(edit_cert_root, text=f"Данные пациента:\n"
                               f"Участок: {data['patient'].get('patient_district')};    "
                               f"№ амб: {data['patient'].get('amb_cart')};\n"
                               f"ФИО: {data['patient'].get('name')};    "
                               f"{data['patient'].get('birth_date')};    "
                               f"пол: {data['patient'].get('gender')};\n"
                               f"Адрес: {data['patient'].get('address')};",

          font=('Comic Sans MS', data.get('text_size')), bg='white').pack(fill='both', expand=True,
                                                                          padx=2, pady=2)
    if not all_data['all_info'].get(type_certificate).get('place_of_requirement'):
        frame_place = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        if type_certificate == 'Оформление в ДДУ / СШ / ВУЗ':
            text = "Куда оформляетесь?:"
            place = ('Детское Дошкольное Учреждение',
                     'Средняя школа (гимназия)',
                     'ВУЗ (колледж)', 'Кадетское училище')
        else:
            text = "Место требования справки:"
            place = all_data.get('place')

        Label(frame_place, text=text,
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        frame_place.columnconfigure(index='all', minsize=40, weight=1)
        frame_place.rowconfigure(index='all', minsize=20)
        frame_place.pack(fill='both', expand=True, padx=2, pady=2)

        destroy_elements['place_of_requirement'] = list()

        def select_place():
            data['certificate']['place_of_requirement'] = selected_place.get()
            for el in destroy_elements.get('place_of_requirement'):
                el.destroy()
            Label(frame_place, text=f" {selected_place.get()}",
                  font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=1, row=0, sticky='e')

            if type_certificate == "Оформление в ДДУ / СШ / ВУЗ" and selected_place.get() == 'ВУЗ (колледж)':
                Label(frame_specialties, text="Специальности для поступления:",
                      font=('Comic Sans MS', data.get('text_size')), bg='white').pack(fill='both', expand=True,
                                                                                      padx=2, pady=2)
                specialties_txt.pack(fill='both', expand=True, padx=2, pady=2)

                frame_specialties.columnconfigure(index='all', minsize=40, weight=1)
                frame_specialties.rowconfigure(index='all', minsize=20)
                frame_specialties.pack(fill='both', after=frame_place, expand=True, padx=2, pady=2)

            frame_place.update()

            # el.quit()

            # frame_place.columnconfigure(index='all', minsize=1, weight=1)
            # frame_place.rowconfigure(index='all', minsize=1)
            # frame_place.config(width=1)

        frame_specialties = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)
        specialties_txt = ScrolledText(frame_specialties, width=80, height=2,
                                       font=('Comic Sans MS', data.get('text_size')), wrap="word")

        selected_place = StringVar()
        frame_place_1 = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)
        destroy_elements['place_of_requirement'].append(frame_place_1)

        row, col = 0, 0
        for mark in place:
            btn = Radiobutton(frame_place_1, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_place, command=select_place)
            btn.grid(row=row, column=col)
            col += 1
            if col == 3:
                row += 1
                col = 0

            frame_place_1.columnconfigure(index='all', minsize=40, weight=1)
            frame_place_1.rowconfigure(index='all', minsize=20)
            frame_place_1.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate in ('На кружки и секции', 'Может работать по специальности...'):

        def close_frame_hobby():
            frame_hobby.destroy()
            frame_hobby.update()

        def append_hobby(event):
            print(event.widget)
            hobby_index = str(event.widget).split('.!')[-1].replace('label', '')
            if hobby_index == '':
                hobby_index = '0'
            hobby = all_data.get('sport_section')[int(hobby_index) - 2]
            if len(hobby_txt.get()) == 0:
                if hobby != '+ СОРЕВНОВАНИЯ':
                    hobby_txt.insert(0, hobby)
                else:
                    hobby_txt.insert(0, 'участия в соревнованиях по ')
            else:
                if hobby != '+ СОРЕВНОВАНИЯ':
                    hobby_txt.insert('end', f", {hobby}")
                else:
                    hobby_txt.insert('end', ' и участия в соревнованиях')

            print('hobby_index', hobby_index)
            print('hobby', hobby)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)
        if type_certificate == 'На кружки и секции':
            txt = 'Может заниматься:'
        else:
            txt = 'Может работать по специальности:'

        Label(frame, text=txt, font=('Comic Sans MS', data.get('text_size')), bg='white').grid(row=0, column=0)
        hobby_txt = Entry(frame, width=70, font=('Comic Sans MS', data.get('text_size')))
        hobby_txt.grid(column=2, row=0, columnspan=3)

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        if type_certificate == 'На кружки и секции':
            frame_hobby = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)
            Label(frame_hobby, text='Кружки и секции',
                  font=('Comic Sans MS', data.get('text_size')), bg='white').grid(row=0, column=0, columnspan=5)

            row, col = 1, 0
            for lbl in all_data.get('sport_section'):
                lbl_0 = Label(frame_hobby, text=lbl,
                              font=('Comic Sans MS', data.get('text_size')), border=1, compound='left',
                              bg='#f0fffe', relief='ridge')
                lbl_0.grid(ipadx=2, ipady=2, padx=2, pady=2, sticky='ew', row=row, column=col)
                lbl_0.bind('<Button-1>', append_hobby)
                col += 1
                if col == 5:
                    col = 0
                    row += 1

            Button(frame_hobby, text='Скрыть', command=close_frame_hobby,
                   font=('Comic Sans MS', data.get('text_size'))).grid(column=col, row=row)

            frame_hobby.columnconfigure(index='all', minsize=40, weight=1)
            frame_hobby.rowconfigure(index='all', minsize=20)
            frame_hobby.pack(fill='both', expand=True, padx=2, pady=2)

            frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

            Label(frame, text="Группа здоровья:",
                  font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

            selected_health_group = StringVar()

            def select_health_group():
                data['certificate']['health_group'] = selected_health_group.get()

            for mark in all_data.get('health').get('group'):
                btn = Radiobutton(frame, text=mark,
                                  font=('Comic Sans MS', data.get('text_size')),
                                  value=mark, variable=selected_health_group, command=select_health_group)
                btn.grid(row=0, column=(all_data.get('health').get('group').index(mark) + 1))

            frame.columnconfigure(index='all', minsize=40, weight=1)
            frame.rowconfigure(index='all', minsize=20)
            frame.pack(fill='both', expand=True, padx=2, pady=2)

            frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

            Label(frame, text="Группа по физ-ре:",
                  font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

            selected_fiz_group = StringVar()

            def select_fiz_group():
                data['certificate']['physical'] = selected_health_group.get()

            for mark in all_data.get('health').get('physical'):
                btn = Radiobutton(frame, text=mark,
                                  font=('Comic Sans MS', data.get('text_size')),
                                  value=mark, variable=selected_fiz_group, command=select_fiz_group)
                btn.grid(row=0, column=(all_data.get('health').get('physical').index(mark) + 1))

            frame.columnconfigure(index='all', minsize=40, weight=1)
            frame.rowconfigure(index='all', minsize=20)
            frame.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate == 'По выздоровлении':
        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Диагноз:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        combo_diagnosis = Combobox(frame, font=('Comic Sans MS', data.get('text_size')), state="readonly")
        combo_diagnosis['values'] = ['ОРИ', "ФРК", "Ветряная оспа"]
        combo_diagnosis.current(0)
        combo_diagnosis.grid(column=1, row=0)

        Label(frame, text="c",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=2, row=0)
        ori_from = Entry(frame, width=15,
                         font=('Comic Sans MS', data.get('text_size')))
        ori_from.grid(column=3, row=0)

        Label(frame, text="по",
              font=('Comic Sans MS', data.get('text_size')), bg='white', compound="center").grid(column=4, row=0)
        ori_until = Entry(frame, width=15,
                          font=('Comic Sans MS', data.get('text_size')))
        ori_until.grid(column=5, row=0)
        ori_until.insert(0, datetime.now().strftime("%d.%m.%Y"))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate in ('Годовой медосмотр',
                            'Оформление в ДДУ / СШ / ВУЗ',
                            'В детский лагерь',
                            'Об усыновлении (удочерении)',
                            'Об отсутствии контактов',
                            'Бесплатное питание',
                            'О нуждаемости в сан-кур лечении'):

        frame_chickenpox = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)
        Label(frame_chickenpox, text="Ветрянка:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)
        chickenpox = ["+", "-", "привит"]
        selected_chickenpox = StringVar()

        def select_chickenpox():
            data['certificate']['chickenpox'] = selected_chickenpox.get()

        for mark in chickenpox:
            btn = Radiobutton(frame_chickenpox, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_chickenpox, command=select_chickenpox)
            btn.grid(row=0, column=(chickenpox.index(mark) + 1))

        Label(frame_chickenpox, text="Аллергия:", font=('Comic Sans MS', data.get('text_size')),
              bg='white').grid(row=1, column=0)
        allergy = ["-", "+"]
        selected_allergy = StringVar()

        allergy_txt = Entry(frame_chickenpox, width=60,
                            font=('Comic Sans MS', data.get('text_size')))

        def select_allergy():
            data['certificate']['allergy'] = selected_allergy.get()
            for el in destroy_elements.get('allergy'):
                el.destroy()
                if selected_allergy.get() == '+':
                    Label(frame_chickenpox, text="Аллергия на:",
                          font=('Comic Sans MS', data.get('text_size')), bg='white').grid(row=1, column=0)
                    allergy_txt.grid(column=1, row=1, columnspan=3)
                else:
                    Label(frame_chickenpox, text="Аллергоанамнез не отягощен",
                          font=('Comic Sans MS', data.get('text_size')),
                          bg='white').grid(row=1, column=1, columnspan=3)

                frame_chickenpox.update()

        destroy_elements['allergy'] = list()
        for mark in allergy:
            btn = Radiobutton(frame_chickenpox, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_allergy, command=select_allergy)
            btn.grid(column=(allergy.index(mark) + 1), row=1)
            destroy_elements['allergy'].append(btn)

        frame_chickenpox.columnconfigure(index='all', minsize=40, weight=1)
        frame_chickenpox.rowconfigure(index='all', minsize=20)
        frame_chickenpox.pack(fill='both', expand=True, padx=2, pady=2)

    def diagnosis_kb():
        local_frame_diagnosis = dict()
        local_destroy_elements = list()

        def select_diagnosis(event):
            print('select_diagnosis event.widget', event.widget)

            frame_ = str(event.widget).split('.!')[2].replace('frame', '')
            label_ = str(event.widget).split('.!')[4].replace('label', '')
            print('select_diagnosis frame_, label_', frame_, label_)

            # for w in str(event.widget).split('.!'):
            #     if 'frame' in w:
            #         frame_ = w.replace('frame', '')
            #     if 'label' in w:
            #         label_ = w.replace('label', '')
            # print(frame_, label_)
            selected_diagnosis = all_data.get('diagnosis')[int(frame_) - 2][int(label_) - 1]
            diagnosis_text.insert(INSERT, f" {selected_diagnosis}")

        def select_category_diagnosis(event):
            widget = ''
            if len(diagnosis_text.get(1.0, 'end')) > 3 and diagnosis_text.get(1.0, 'end')[-3:-1] != '. ':
                diagnosis_text.insert(INSERT, ". ")

            for w in str(event.widget).split('.!'):
                if 'frame' in w:
                    widget = w.replace('frame', '')
            open_button = all_data.get('diagnosis')[int(widget) - 2][0]
            for el in local_destroy_elements:
                el.destroy()
            local_destroy_elements.clear()
            print(open_button)
            print('local_destroy_elements', local_destroy_elements)
            print('local_frame_diagnosis', local_frame_diagnosis)
            for key_diagnosis in local_frame_diagnosis:
                if key_diagnosis != open_button:
                    lbl_dig = Label(master=local_frame_diagnosis.get(key_diagnosis), text=f"{key_diagnosis}",
                                    font=('Comic Sans MS', data.get('text_size')), bg='white')
                    lbl_dig.grid(column=0, row=0)
                    lbl_dig.bind('<Button-1>', select_category_diagnosis)
                    local_destroy_elements.append(lbl_dig)

            frame_diagnosis_in = Frame(master=local_frame_diagnosis.get(open_button),
                                       borderwidth=1, relief="solid", padx=4, pady=4)
            local_destroy_elements.append(frame_diagnosis_in)

            lbl_dig = Label(frame_diagnosis_in, text=f"{all_data.get('diagnosis')[int(widget) - 2][0]}",
                            font=('Comic Sans MS', data.get('text_size')), bg='white')
            lbl_dig.grid(column=0, row=0, columnspan=5)
            lbl_dig.bind('<Button-1>', select_category_diagnosis)

            row_, col_ = 1, 0
            for lbl_dig in all_data.get('diagnosis')[int(widget) - 2][1:]:
                if col_ == 5:
                    row_ += 1
                    col_ = 0
                lbl_01 = Label(frame_diagnosis_in, text=lbl_dig,
                               font=('Comic Sans MS', data.get('text_size')), border=1,
                               compound='left',
                               bg='#f0fffe', relief='ridge')
                lbl_01.grid(ipadx=2, ipady=2, padx=2, pady=2, column=col_, row=row_)
                lbl_01.bind('<Button-1>', select_diagnosis)
                col_ += 1

            frame_diagnosis_in.columnconfigure(index='all', minsize=40, weight=1)
            frame_diagnosis_in.rowconfigure(index='all', minsize=20)
            frame_diagnosis_in.pack(fill='both', expand=True, padx=2, pady=2)

            # frame_diagnosis_in.grid(column=0, row=0)

            # if tuple_diagnosis[0] == diagnosis_local_data.get('open_buttons', ''):

            # diagnosis_root.destroy()
            # diagnosis_kb()

        def close_diagnosis_kb():
            diagnosis.delete(1.0, 'end')
            diagnosis.insert(INSERT, diagnosis_text.get(1.0, 'end')[:-1])

            diagnosis.focus()
            diagnosis_root.destroy()

        diagnosis_root = Toplevel()
        diagnosis_root.title('Клавиатура диагнозов')
        diagnosis_root.config(bg='white')

        frame_diagnosis = Frame(diagnosis_root, borderwidth=1, relief="solid", padx=2, pady=2)
        Label(frame_diagnosis, text="Диагноз:", font=('Comic Sans MS', 15), bg='white').grid()
        diagnosis_text = ScrolledText(frame_diagnosis, width=70, height=10, font=('Comic Sans MS', 15),
                                      wrap="word")
        if diagnosis.get(1.0, 'end') == ' \n' or diagnosis.get(1.0, 'end') == '\n':
            diagnosis.delete(1.0, 'end')
        elif len(diagnosis.get(1.0, 'end')) > 1 and diagnosis.get(1.0, 'end') == '\n':
            if diagnosis.get(1.0, 'end') == '\n':
                diagnosis_text.insert(INSERT, diagnosis.get(1.0, 'end')[:-1])
        else:
            diagnosis_text.insert(INSERT, diagnosis.get(1.0, 'end'))
        diagnosis_text.focus()
        diagnosis_text.grid(column=0, row=1, rowspan=6)
        Button(frame_diagnosis, text='Закрыть\nклавиатуру',
               command=close_diagnosis_kb,
               font=('Comic Sans MS', data.get('text_size'))).grid(column=1, row=3)
        frame_diagnosis.columnconfigure(index='all', minsize=40, weight=1)
        frame_diagnosis.rowconfigure(index='all', minsize=20)
        frame_diagnosis.pack(fill='both', expand=True, padx=2, pady=2)

        # frame_diagnosis.grid(padx=2, pady=2, column=0, row=0)

        for tuple_diagnosis in all_data.get('diagnosis'):
            frame_diagnosis = Frame(diagnosis_root, borderwidth=1, relief="solid", padx=4, pady=4)

            # local_destroy_elements[tuple_diagnosis[0]] = list()
            local_frame_diagnosis[tuple_diagnosis[0]] = frame_diagnosis

            lbl_d = Label(frame_diagnosis, text=f"{tuple_diagnosis[0]}",
                          font=('Comic Sans MS', data.get('text_size')), bg='white')
            local_destroy_elements.append(lbl_d)
            lbl_d.grid(column=0, row=0)
            lbl_d.bind('<Button-1>', select_category_diagnosis)

            frame_diagnosis.columnconfigure(index='all', minsize=40, weight=1)
            frame_diagnosis.rowconfigure(index='all', minsize=20)
            frame_diagnosis.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate == 'Оформление в ДДУ / СШ / ВУЗ':
        frame_injury_operation = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame_injury_operation, text="Травмы и операции:", font=('Comic Sans MS', data.get('text_size')),
              bg='white').grid(row=0, column=0)
        injury_operation = ("-", "+")
        selected_injury_operation = StringVar()
        injury_operation_txt = Entry(frame_injury_operation, width=60,
                                     font=('Comic Sans MS', data.get('text_size')))

        def select_injury_operation():
            injury_operation_val = selected_injury_operation.get()
            data['certificate']['injury_operation'] = injury_operation_val
            for el in destroy_elements.get('injury_operation'):
                el.destroy()

            if injury_operation_val == '-':
                Label(frame_injury_operation, text="не было",
                      font=('Comic Sans MS', data.get('text_size')), bg='white').grid(row=0, column=1)
            else:
                injury_operation_txt.grid(column=1, row=0)

        destroy_elements['injury_operation'] = list()
        for mark in injury_operation:
            btn = Radiobutton(frame_injury_operation, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_injury_operation, command=select_injury_operation)
            btn.grid(column=(injury_operation.index(mark) + 1), row=0)
            destroy_elements['injury_operation'].append(btn)

        frame_injury_operation.columnconfigure(index='all', minsize=40, weight=1)
        frame_injury_operation.rowconfigure(index='all', minsize=20)
        frame_injury_operation.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate in ('Годовой медосмотр',
                            'Оформление в ДДУ / СШ / ВУЗ',
                            'В детский лагерь',
                            'Об усыновлении (удочерении)'):

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Рост (см):",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)
        height = Entry(frame, width=15,
                       font=('Comic Sans MS', data.get('text_size')))
        height.grid(column=1, row=0)

        Label(frame, text="    Вес (кг):",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=2, row=0)
        weight = Entry(frame, width=15,
                       font=('Comic Sans MS', data.get('text_size')))
        weight.grid(column=3, row=0)

        vision = Entry(frame, width=15,
                       font=('Comic Sans MS', data.get('text_size')))
        if type_certificate in ('Годовой медосмотр', 'Оформление в ДДУ / СШ / ВУЗ'):
            Label(frame, text="    Зрение:",
                  font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=4, row=0)
            vision.grid(column=5, row=0)
            age = get_age(data['patient'].get('birth_date'))
            if age >= 5:
                vision.insert(0, '1.0/1.0')
            else:
                vision.insert(0, 'предметное')

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Диагноз:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)
        diagnosis = ScrolledText(frame, width=80, height=4,
                                 font=('Comic Sans MS', data.get('text_size')), wrap="word")
        diagnosis.grid(column=0, row=1, rowspan=4, columnspan=4)

        selected_health_group = StringVar()
        selected_fiz_group = StringVar()
        regime_vars = dict()
        for mark in all_data.get('health').get('regime'):
            regime_vars[mark] = IntVar()
        selected_diet = StringVar()
        desk_vars = dict()
        for mark in all_data.get('health').get('desk'):
            desk_vars[mark] = IntVar()

        def diagnosis_healthy():
            diagnosis.delete(1.0, 'end')
            if data['patient'].get('gender', '') == 'женский':
                diagnosis.insert(INSERT, 'Соматически здорова. ')
            else:
                diagnosis.insert(INSERT, 'Соматически здоров. ')
            selected_health_group.set('1')
            data['certificate']['health_group'] = selected_health_group.get()
            selected_fiz_group.set('Основная')
            data['certificate']['physical'] = selected_health_group.get()
            regime_vars['общий'].set(1)
            data['certificate']['regime'] = ["общий"]
            selected_diet.set('Б')
            data['certificate']['diet'] = selected_health_group.get()
            desk_vars['по росту'].set(1)
            data['certificate']['desk'] = ["по росту"]

        Button(frame, text='Здоров', command=diagnosis_healthy,
               font=('Comic Sans MS', data.get('text_size'))).grid(column=1, row=0)

        Button(frame, text='Клавиатура диагнозов', command=diagnosis_kb,
               font=('Comic Sans MS', data.get('text_size'))).grid(column=2, row=0)

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Группа здоровья:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        def select_health_group():
            data['certificate']['health_group'] = selected_health_group.get()

        for mark in all_data.get('health').get('group'):
            btn = Radiobutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_health_group, command=select_health_group)
            btn.grid(row=0, column=(all_data.get('health').get('group').index(mark) + 1))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Группа по физ-ре:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        def select_fiz_group():
            data['certificate']['physical'] = selected_health_group.get()

        for mark in all_data.get('health').get('physical'):
            btn = Radiobutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_fiz_group, command=select_fiz_group)
            btn.grid(row=0, column=(all_data.get('health').get('physical').index(mark) + 1))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Режим:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        def select_regime():
            result = list()
            for regime in all_data.get('health').get('regime'):
                if regime_vars.get(regime).get() == 1:
                    result.append(regime)
            data['certificate']['regime'] = result
            print(result)


        for mark in all_data.get('health').get('regime'):
            btn = Checkbutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              variable=regime_vars.get(mark), command=select_regime,
                              onvalue=1, offvalue=0)
            btn.grid(row=0, column=(all_data.get('health').get('regime').index(mark) + 1))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Стол:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        def select_diet():
            data['certificate']['diet'] = selected_health_group.get()

        for mark in all_data.get('health').get('diet'):
            btn = Radiobutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              value=mark, variable=selected_diet, command=select_diet)
            btn.grid(row=0, column=(all_data.get('health').get('diet').index(mark) + 1))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Парта:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)

        def select_desk():
            result = list()
            for desk in all_data.get('health').get('desk'):
                if desk_vars.get(desk).get() == 1:
                    result.append(desk)
            data['certificate']['desk'] = result
            print(result)

        for mark in all_data.get('health').get('desk'):
            btn = Checkbutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              variable=desk_vars.get(mark), command=select_desk,
                              onvalue=1, offvalue=0)
            btn.grid(row=0, column=(all_data.get('health').get('desk').index(mark) + 1))

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

    if type_certificate == 'О нуждаемости в сан-кур лечении':
        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Диагноз:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0)
        diagnosis = ScrolledText(frame, width=80, height=4,
                                 font=('Comic Sans MS', data.get('text_size')), wrap="word")
        diagnosis.grid(column=0, row=1, rowspan=4, columnspan=4)

        def diagnosis_healthy():
            diagnosis.delete(1.0, 'end')
            if data['patient'].get('gender', '') == 'женский':
                diagnosis.insert(INSERT, 'Соматически здорова. ')
            else:
                diagnosis.insert(INSERT, 'Соматически здоров. ')

        Button(frame, text='Здоров', command=diagnosis_healthy,
               font=('Comic Sans MS', data.get('text_size'))).grid(column=1, row=0)

        Button(frame, text='Клавиатура диагнозов', command=diagnosis_kb,
               font=('Comic Sans MS', data.get('text_size'))).grid(column=2, row=0)

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

        frame = Frame(edit_cert_root, borderwidth=1, relief="solid", padx=4, pady=4)

        Label(frame, text="Профиль санатория:",
              font=('Comic Sans MS', data.get('text_size')), bg='white').grid(column=0, row=0, columnspan=3)

        def select_profile():
            result = list()
            for regime in sanatorium_profile:
                if sanatorium_profile.get(regime).get() == 1:
                    result.append(regime)
            data['certificate']['sanatorium_profile'] = result
            print(result)

        sanatorium_profile = dict()
        for profile in ('пульмонологического', 'гастроэнтерологического', 'ревматологического',
                        'неврологического', 'эндокринологического', 'нефрологического',
                        'гинекологического', 'кардиологического', 'дерматологического',
                        'ортопедотравматологического', 'офтальмологического'):
            sanatorium_profile[profile] = IntVar()
        row, col = 1, 0
        for mark in sanatorium_profile:
            btn = Checkbutton(frame, text=mark,
                              font=('Comic Sans MS', data.get('text_size')),
                              variable=sanatorium_profile.get(mark), command=select_profile,
                              onvalue=1, offvalue=0)
            btn.grid(row=row, column=col)
            col += 1
            if col == 3:
                col = 0
                row += 1

        frame.columnconfigure(index='all', minsize=40, weight=1)
        frame.rowconfigure(index='all', minsize=20)
        frame.pack(fill='both', expand=True, padx=2, pady=2)

    def create_certificate():
        try:
            for marker in all_data['all_info'].get(type_certificate):
                render_data[marker] = all_data['all_info'][type_certificate].get(marker)

            render_data['time'] = datetime.now().strftime("%H:%M")
            render_data['number_cert'] = ''
            render_data['name'] = data['patient'].get('name')
            render_data['birth_date'] = data['patient'].get('birth_date')
            render_data['gender'] = data['patient'].get('gender')
            render_data['address'] = data['patient'].get('address')
            render_data['amb_cart'] = data['patient'].get('amb_cart')

            if not render_data.get('place_of_requirement'):
                render_data['place_of_requirement'] = data['certificate'].get('place_of_requirement')

            render_data['type'] = type_certificate
            if type_certificate == 'Оформление в ДДУ / СШ / ВУЗ':
                place_of_req = render_data.get('place_of_requirement')
                if place_of_req == 'Детское Дошкольное Учреждение':
                    render_data['type'] = 'Оформление в Детское Дошкольное Учреждение'
                    render_data['recommendation'] = \
                        render_data.get('recommendation').replace('Режим _',
                                                                  'Режим щадящий 1 мес, затем Режим _') \
                        + "\nМебель по росту"
                if place_of_req == 'Средняя школа (гимназия)':
                    render_data['place_of_requirement'] = 'Средняя школа (гимназия)'
                    render_data['type'] = 'Оформление в Среднюю школу (гимназию)'
                    if data['patient'].get('gender', '') == 'женский':
                        render_data['diagnosis'] += '\nГотова к обучению в ' \
                                                    'общеобразовательной школе с _______ лет'
                    else:
                        render_data['diagnosis'] += '\nГотов к обучению в ' \
                                                    'общеобразовательной школе с _______ лет'
                    data['certificate']['recommendation'] += " Парта _"

                if place_of_req == 'ВУЗ (колледж)':
                    render_data['type'] = 'Оформление в ВУЗ'
                    render_data['additional_medical_information'] = \
                        render_data.get('additional_medical_information').replace(' Vis OD/OS = __________ ;', '')
                    render_data['recommendation'] = "_____________________________________________"
                    render_data['place_of_requirement'] = 'Для поступления в учреждения высшего, ' \
                                                          'среднего специального и ' \
                                                          'профессионально-технического образования '
                    render_data['diagnosis'] += "\nОтсутствуют медицинские противопоказания " \
                                                f"к обучению по специальности: \n" \
                                                f"{specialties_txt.get(1.0, 'end')}" \
                                                "(пункт 2 приложения к постановлению " \
                                                "МЗ РБ от 25.07.2022г. №71)"

                if place_of_req == 'Кадетское училище':
                    render_data['type'] = 'Оформление в Кадетское училище'
                    render_data['place_of_requirement'] = 'Для обучения в кадетском училище '

                    render_data['additional_medical_information'] = render_data.get('additional_medical_information')\
                        + '\nОфтальмолог: ________________________________________________________' \
                          '\nНевролог: ___________________________________________________________' \
                          '\nОториноларинголог : _________________________________________________' \
                          '\nСтоматолог: _________________________________________________________' \
                          '\nХирург: _____________________________________________________________' \
                          '\nПедиатр: ____________________________________________________________' \
                          '\nОАК: ________________________________________________________________' \
                          '\nОАМ: ________________________________________________________________' \
                          '\nУЗИ сердца: _________________________________________________________' \
                          '\nУЗИ щитовидной  железы : ____________________________________________'

                    render_data['diagnosis'] += '\nФизическое развитие (выше- ниже-) среднее, (дис-) гармоничное\n' \
                                                'Врачебное профессионально-консультативное заключение: ' \
                                                'отсутствуют медицинские противопоказания к обучению  в ГУО ' \
                                                '«Минском городском  кадетском училище».'
                    render_data['recommendation'] = f"{render_data.get('recommendation')} Парта _"

            if render_data.get('date_of_issue') == 'now':
                render_data['date_of_issue'] = datetime.now().strftime("%d.%m.%Y")

            if type_certificate == 'По выздоровлении':
                render_data['diagnosis'] = f"{combo_diagnosis.get()} c {ori_from.get()} по {ori_until.get()}"
                data['certificate']['date_of_issue'] = ori_until.get()

            if type_certificate == 'Годовой медосмотр':
                date = data['patient'].get('birth_date')
                while datetime.now() > datetime.strptime(date, "%d.%m.%Y"):
                    day, month, year = date.split('.')
                    year = str(int(year) + 1)
                    date = '.'.join([day, month, year])
                if (datetime.now().month == datetime.strptime(date, '%d.%m.%Y').month or
                    (datetime.now() + timedelta(30)).month == datetime.strptime(date, "%d.%m.%Y").month) and \
                        datetime.now().year >= datetime.strptime(date, '%d.%m.%Y').year:
                    day, month, year = date.split('.')
                    year = str(int(year) + 1)
                    date = '.'.join([day, month, year])

                render_data['validity_period'] = (datetime.strptime(date, '%d.%m.%Y') -
                                                  timedelta(1)).strftime('%d.%m.%Y')

            if type_certificate == 'На кружки и секции':
                render_data['place_of_requirement'] = \
                    f"{render_data.get('place_of_requirement')}{hobby_txt.get()}" \
                    f"".replace('участия в соревнованиях по ', '').replace(' и участия в соревнованиях', '')

                render_data['diagnosis'] = \
                    f"{render_data.get('diagnosis')}{hobby_txt.get()}"

                if data['certificate'].get('health_group'):
                    render_data['diagnosis'] = \
                        f"{render_data.get('diagnosis')}\n " \
                        f"Группа здоровья: {data['certificate'].get('health_group')};"

                if data['certificate'].get('physical'):
                    render_data['diagnosis'] = \
                        f"{render_data.get('diagnosis')}" \
                        f"Группа по физкультуре: {data['certificate'].get('physical')};"

            if type_certificate in ('Годовой медосмотр',
                                    'Оформление в ДДУ / СШ / ВУЗ',
                                    'В детский лагерь',
                                    'Об усыновлении (удочерении)',
                                    'Об отсутствии контактов',
                                    'Бесплатное питание',
                                    'О нуждаемости в сан-кур лечении'):
                for ex_marker in ('chickenpox', 'allergy'):
                    if not data['certificate'].get(ex_marker):
                        if ex_marker == 'chickenpox':
                            messagebox.showinfo('Ошибка!', 'Не указана ветрянка!')
                        if ex_marker == 'allergy':
                            messagebox.showinfo('Ошибка!', 'Не указана аллергия!')
                        raise ValueError
                text_past_illnesses = ''
                render_data['chickenpox'] = data['certificate'].get('chickenpox')
                if data['certificate'].get('chickenpox', '') == '+':
                    text_past_illnesses += 'ОРИ, Ветряная оспа; '
                if data['certificate'].get('chickenpox', '') == '-':
                    text_past_illnesses += 'ОРИ; Ветряной оспой не болел; '
                if data['certificate'].get('chickenpox', '') == 'привит':
                    text_past_illnesses += 'ОРИ; от ветряной оспы привит; '

                render_data['allergy'] = data['certificate'].get('allergy', '')
                if data['certificate'].get('allergy', '') == '-':
                    text_past_illnesses += 'Аллергоанамнез: не отягощен; '
                if data['certificate'].get('allergy', '') == '+':
                    text_past_illnesses += 'Аллергоанамнез отягощен: '

                    if len(allergy_txt.get()) > 1:
                        text_past_illnesses += f'\nАллергия на: {allergy_txt.get()}'
                        data['certificate']['allergy'] = f"+\n{allergy_txt.get()}"
                        render_data['allergy'] = f"{render_data.get('allergy')}\n{allergy_txt.get()}"
                "injury_operation"
                if type_certificate == 'Оформление в ДДУ / СШ / ВУЗ':

                    render_data['injury'] = data['certificate'].get('injury_operation', '')
                    if data['certificate'].get('injury_operation', '') == '-':
                        text_past_illnesses += 'Травм и операций не было; '
                    if data['certificate'].get('allergy', '') == '+':
                        text_past_illnesses += 'Травмы и операции: '
                        if len(injury_operation_txt.get()) > 1:
                            text_past_illnesses += f'\n{injury_operation_txt.get()}'
                            data['certificate']['injury'] = f"+\n{injury_operation_txt.get()}"
                            render_data['injury'] = f"{render_data.get('allergy')}\n{allergy_txt.get()}"

                if render_data.get('past_illnesses'):
                    render_data['past_illnesses'] = f"{text_past_illnesses}\n{render_data.get('past_illnesses')}"
                else:
                    render_data['past_illnesses'] = text_past_illnesses

            if type_certificate in ('Годовой медосмотр',
                                    'Оформление в ДДУ / СШ / ВУЗ',
                                    'В детский лагерь',
                                    'Об усыновлении (удочерении)'):

                for ex_marker in ('health_group', 'physical', 'regime', 'diet', 'desk'):
                    if not data['certificate'].get(ex_marker):
                        if ex_marker == 'health_group':
                            messagebox.showinfo('Ошибка!', 'Не указана группа здоровья!')
                        if ex_marker == 'physical':
                            messagebox.showinfo('Ошибка!', 'Не указана группа по физкультуре!')
                        if ex_marker == 'regime':
                            messagebox.showinfo('Ошибка!', 'Не указан режим!')
                        if ex_marker == 'diet':
                            messagebox.showinfo('Ошибка!', 'Не указана диета!')
                        if ex_marker == 'desk':
                            messagebox.showinfo('Ошибка!', 'Не указана рассадка!')
                        raise ValueError

                if not weight.get():
                    messagebox.showinfo('Ошибка!', 'Не указан вес!')
                    raise ValueError

                if not weight.get().replace('.', '').replace(',', '').isdigit():
                    messagebox.showinfo('Ошибка!', 'Укажите вес цифрами!')
                    raise ValueError

                if not height.get():
                    messagebox.showinfo('Ошибка!', 'Не указан рост!')
                    raise ValueError

                if not height.get().replace('.', '').replace(',', '').isdigit():
                    messagebox.showinfo('Ошибка!', 'Укажите рост цифрами!')
                    raise ValueError

                if type_certificate in ('Годовой медосмотр', 'Оформление в ДДУ / СШ / ВУЗ'):
                    render_data['visus'] = f"VIS OD/OS\n= {vision.get()}\n"
                    render_data['additional_medical_information'] = \
                        render_data.get('additional_medical_information',
                                        '').replace('Vis OD/OS = __________', f"Vis OD/OS = {vision.get()}")
                else:
                    render_data['visus'] = ''

                render_data['height'] = height.get().replace(',', '.')
                render_data['weight'] = weight.get().replace(',', '.')
                render_data['group'] = selected_health_group.get()
                render_data['physical'] = selected_fiz_group.get()

                print('height', f"'{height.get()}'")
                print('weight', f"'{weight.get()}'")

                add_med_info = render_data.get('additional_medical_information', '')
                add_med_info = add_med_info.replace('Рост _____ см', f'Рост {height.get()} см')
                add_med_info = add_med_info.replace('Вес _____ кг', f'Вес {weight.get()} кг')
                render_data['additional_medical_information'] = add_med_info

                diagnosis_certificate = render_data.get('diagnosis', '')
                if len(diagnosis.get(1.0, 'end')) > 2:
                    if diagnosis.get(1.0, 'end').endswith('\n'):
                        diagnosis_certificate = f"{diagnosis.get(1.0, 'end')[:-2]}\n{diagnosis_certificate}"
                    else:
                        diagnosis_certificate = f"{diagnosis.get(1.0, 'end')}\n{diagnosis_certificate}"

                diagnosis_certificate = diagnosis_certificate.replace('Группа здоровья: _',
                                                                      f'Группа здоровья: {selected_health_group.get()}')
                diagnosis_certificate = diagnosis_certificate.replace('Группа по физкультуре: _',
                                                                      f'Группа по физкультуре: {selected_fiz_group.get()}')
                render_data['diagnosis'] = diagnosis_certificate

                recommendation = render_data.get('recommendation', '')
                result = ''
                for regime in data['certificate'].get('regime'):
                    result += f"{regime}, "
                result = result[:-2]
                recommendation = recommendation.replace('Режим _', f'Режим {result}')
                render_data['regime'] = result

                recommendation = recommendation.replace('Стол _', f'Стол {selected_diet.get()}')
                render_data['diet'] = selected_diet.get()

                if type_certificate in ('Годовой медосмотр',
                                        'Оформление в ДДУ / СШ / ВУЗ'):
                    result = ''
                    desk_num = list()
                    for desk in data['certificate'].get('desk'):
                        if desk.isdigit():
                            desk_num.append(int(desk))
                    for desk in sorted(desk_num):
                        result += f'{desk} - '
                    if result:
                        result = result[:-2]
                    if 'средний ряд' in data['certificate'].get('desk'):
                        result += 'средний ряд '
                    if 'по росту' in data['certificate'].get('desk'):
                        result += 'по росту'
                    if type_certificate in ('Годовой медосмотр', 'Оформление в ДДУ / СШ / ВУЗ'):
                        if 'Детское Дошкольное Учреждение' in data['certificate'].get('place_of_requirement'):
                            recommendation = recommendation.replace('Парта _', f" Мебель {result}")
                            render_data['desk'] = f"Мебель {result}"
                        else:
                            recommendation = recommendation.replace('Парта _', f" Парта {result}")
                            render_data['desk'] = f"Парта {result}"

                render_data['recommendation'] = recommendation

            if type_certificate == 'О нуждаемости в сан-кур лечении':
                diagnosis_certificate = render_data.get('diagnosis', '')
                if len(diagnosis.get(1.0, 'end')) > 2:
                    if diagnosis.get(1.0, 'end').endswith('\n'):
                        diagnosis_certificate = f"{diagnosis.get(1.0, 'end')[:-2]}\n{diagnosis_certificate}"
                    else:
                        diagnosis_certificate = f"{diagnosis.get(1.0, 'end')}\n{diagnosis_certificate}"
                render_data['diagnosis'] = diagnosis_certificate

                profile_rec = 'Ребенок нуждается в санаторно-курортном лечении: \n'
                for i in data['certificate'].get('sanatorium_profile'):
                    profile_rec += f"{i} профиля \n"
                render_data['recommendation'] = profile_rec[:-2]

            for key, value in render_data.items():
                print(key, value)
        except ValueError:
            pass
        else:
            create_doc()

    Button(edit_cert_root, text='Создать справку', command=create_certificate,
           font=('Comic Sans MS', data.get('text_size'))).pack(fill='both', expand=True, padx=2, pady=2)

    edit_cert_root.mainloop()


def get_age(birth_date):
    birthday = ''
    for i in birth_date:
        if i.isdigit():
            birthday += i
        else:
            birthday += '.'
    birthday = birthday.split('.')
    date_form = "%d.%m.%Y"
    if len(birthday[-1]) == 2:
        date_form = "%d.%m.%y"
    birthday = '.'.join(birthday)
    birthday = datetime.strptime(birthday, date_form)

    today = datetime.today()
    age = today.year - birthday.year

    if (today.month < birthday.month) or (today.month == birthday.month and today.day < birthday.day):
        age -= 1

    return age


def create_doc():
    type_certificate = data['certificate'].get('type_certificate', '')

    if not render_data.get('additional_medical_information'):
        render_data['additional_medical_information'] = '_' * 60
    if not render_data.get('past_illnesses'):
        render_data['past_illnesses'] = '_' * 60
    if not render_data.get('recommendation'):
        render_data['recommendation'] = '_' * 50

    render_data['doctor_name'] = data['doctor'].get('doctor_name')

    if (type_certificate in ('Годовой медосмотр', 'Оформление в ДДУ / СШ / ВУЗ')
            and render_data.get('place_of_requirement') in ('Средняя школа (гимназия)',
                                                            'Детское Дошкольное Учреждение')):
        render_data['recommendation'] = f"{render_data.get('recommendation', '')} \nРазрешены занятия в бассейне"
        if (data['doctor'].get('doctor_district') == '19'
                and render_data.get('place_of_requirement') != 'Средняя школа (гимназия)'):
            render_data['recommendation'] = \
                render_data.get('recommendation', '').replace('\nРазрешены занятия в бассейне', '')

    if type_certificate in ('ЦКРОиР', 'О нуждаемости в сан-кур лечении',
                            'Об усыновлении (удочерении)', 'Бесплатное питание') \
            or (data['certificate'].get('type_certificate') == 'Оформление в ДДУ / СШ / ВУЗ' and not
    ('Для поступления в учреждения высшего' in render_data.get('place_of_requirement') or
     ('Для обучения в кадетском училище' in render_data.get('place_of_requirement')))):
        doctor_name, district, pediatric_division = (data['doctor'].get('doctor_name'),
                                                     data['doctor'].get('doctor_district'),
                                                     data['doctor'].get('ped_div'))
        if pediatric_division in ('1', '2'):
            if data['certificate'].get('type_certificate') in ('ЦКРОиР', 'Бесплатное питание'):
                type_cert = '7.9'
            else:
                type_cert = '7.6'
            info = [pediatric_division,
                    district,
                    None,
                    datetime.now().strftime("%d.%m.%Y"),
                    render_data.get('name'),
                    render_data.get('birth_date'),
                    render_data.get('address'),
                    type_cert,
                    doctor_name]
            number = data_base.save_certificate_ped_div(data=info,
                                                        type_table='certificate_ped_div',
                                                        district_pd=pediatric_division)
            render_data['number_cert'] = f"№ {number}"
        else:
            render_data['number_cert'] = f"№ _________"

    if data['certificate'].get('type_certificate') in ('Годовой медосмотр',
                                                       'Оформление в ДДУ / СШ / ВУЗ',
                                                       'В детский лагерь',
                                                       'Об усыновлении (удочерении)'):

        age = get_age(data['patient'].get('birth_date'))
        if age >= 5:
            render_data['covid_vac'] = 'Предложена вакцинация против инфекции COVID-19\n' \
                                       'Родители отказываются от проведения вакцинации'
        else:
            render_data['covid_vac'] = ' '

        if data['certificate'].get('type_certificate') == 'Оформление в ДДУ / СШ / ВУЗ' or age >= 11:
            render_data['hearing'] = '\nСлух в норме.'
        else:
            render_data['hearing'] = ''

        if age >= 4:
            render_data['posture'] = '\nОсанка не нарушена.'
        else:
            render_data['posture'] = ''
        anthropometry = {
            'мужской': {
                'weight': {
                    1: [8.5, 8.9, 9.4, 10, 10.9, 11.6, 12.1, 1000],
                    2: [10.6, 11, 11.7, 12.6, 13.5, 14.2, 15, 1000],
                    3: [12.1, 12.8, 13.8, 14.8, 16, 16.9, 17.7, 1000],
                    4: [13.4, 14.2, 15.1, 16.4, 17.8, 19.4, 20.3, 1000],
                    5: [14.8, 15.7, 16.8, 18.3, 20, 21.7, 23.4, 1000],
                    6: [16.3, 17.5, 18.8, 20.4, 22.6, 24.7, 26.7, 1000],
                    7: [18, 19.5, 21, 22.9, 25.4, 28, 30.8, 1000],
                    8: [20, 21.5, 23.3, 25.5, 28.3, 31.4, 35.5, 1000],
                    9: [21.9, 23.5, 25.6, 28.1, 31.5, 35.1, 39.1, 1000],
                    10: [23.9, 25.6, 28.2, 31.4, 35.1, 39.7, 44.7, 1000],
                    11: [26, 28, 31, 34.9, 39.9, 44.9, 51.5, 1000],
                    12: [28.2, 30.7, 34.4, 38.8, 45.1, 50.6, 58.7, 1000],
                    13: [30.9, 33.8, 38, 43.4, 50.6, 56.8, 66, 1000],
                    14: [34.3, 38, 42.8, 48.8, 56.6, 63.4, 73.2, 1000],
                    15: [38.7, 43, 48.3, 54.8, 62.8, 70, 80.1, 1000],
                    16: [44, 48.3, 54, 61, 69.6, 76.5, 84.7, 1000],
                    17: [49.3, 54.6, 59.8, 66.3, 74, 80.1, 87.8, 1000]},
                'height': {
                    1: [71.2, 72.3, 74, 75.5, 77.3, 79.7, 81.7, 1000],
                    2: [81.3, 83, 84.5, 86.8, 89, 90.8, 94, 1000],
                    3: [88, 90, 92.3, 96, 99.8, 102, 104.5, 1000],
                    4: [93.2, 95.5, 98.3, 102, 105.5, 108, 110.6, 1000],
                    5: [98.9, 101.5, 104.4, 108.3, 112, 114.5, 117, 1000],
                    6: [105, 107.7, 110.9, 115, 118.7, 121.1, 123.8, 1000],
                    7: [111, 113.6, 116.8, 121.2, 125, 128, 130.6, 1000],
                    8: [116.3, 119, 122.1, 126.9, 130.8, 134.5, 137, 1000],
                    9: [121.5, 124.7, 125.6, 133.4, 136.3, 140.3, 143, 1000],
                    10: [126.3, 129.4, 133, 137.8, 142, 146.7, 149.2, 1000],
                    11: [131.3, 134.5, 138.5, 143.2, 148.3, 152.9, 156.2, 1000],
                    12: [136.2, 140, 143.6, 149.2, 154.5, 159.5, 163.5, 1000],
                    13: [141.8, 145.7, 149.8, 154.8, 160.6, 166, 170.7, 1000],
                    14: [148.3, 152.3, 156.2, 161.2, 167.7, 172, 176.7, 1000],
                    15: [154.6, 158.6, 162.5, 166.8, 173.5, 177.6, 181.6, 1000],
                    16: [158.8, 163.2, 166.8, 173.3, 177.8, 182, 186.3, 1000],
                    17: [162.8, 166.6, 171.6, 177.3, 181.6, 186, 188.5, 1000]}},

            'женский': {
                'weight': {
                    1: [8, 8.5, 9, 9.6, 10.2, 10.8, 11.3, 1000],
                    2: [10.2, 10.8, 11.3, 12.1, 12.8, 13.5, 14.1, 1000],
                    3: [11.7, 12.5, 13.3, 13.7, 15.5, 16.5, 17.6, 1000],
                    4: [13, 14, 14.8, 15.9, 17.6, 18.9, 20, 1000],
                    5: [14.7, 15.7, 16.6, 18.1, 19.7, 21.6, 23.2, 1000],
                    6: [16.3, 17.4, 18.7, 20.4, 22.5, 24.8, 27.1, 1000],
                    7: [17.9, 19.4, 20.6, 22.7, 25.3, 28.3, 31.6, 1000],
                    8: [20, 21.4, 23, 25.1, 28.5, 32.1, 36.3, 1000],
                    9: [21.9, 23.4, 25.5, 28.2, 32, 36.3, 41, 1000],
                    10: [22.7, 25, 27.7, 30.6, 34.9, 39.8, 47.4, 1000],
                    11: [24.9, 27.8, 30.7, 34.3, 38.9, 44.6, 55.2, 1000],
                    12: [27.8, 31.8, 36, 40, 45.4, 51.8, 63.4, 1000],
                    13: [32, 38.7, 43, 47.5, 52.5, 59, 69, 1000],
                    14: [37.6, 43.8, 48.2, 52.8, 58, 64, 72.2, 1000],
                    15: [42, 46.8, 50.6, 55.2, 60.4, 66.5, 74.9, 1000],
                    16: [45.2, 48.4, 51.8, 56.5, 61.3, 67.6, 75.6, 1000],
                    17: [46.2, 49.2, 52.9, 57.3, 61.9, 68, 76, 1000]},
                'height': {
                    1: [70.1, 71.4, 72.8, 74.1, 75.8, 78, 79.6, 1000],
                    2: [80.1, 81.7, 83.3, 85.2, 87.5, 90.1, 92.5, 1000],
                    3: [89, 90.8, 93, 95.5, 98.1, 100.7, 103.1, 1000],
                    4: [94, 96.1, 98.5, 101.5, 104.1, 106.9, 109.7, 1000],
                    5: [100, 102.5, 104.7, 107.5, 110.7, 113.6, 116.7, 1000],
                    6: [105.3, 108, 110.9, 114.1, 118, 120.6, 124, 1000],
                    7: [111.1, 113.6, 116.9, 120.8, 124.8, 128, 131.3, 1000],
                    8: [116.5, 119.3, 123, 127.2, 131, 134.3, 137.7, 1000],
                    9: [122, 124.8, 128.4, 132.8, 137, 140.5, 144.8, 1000],
                    10: [127, 130.5, 134.3, 139, 142.9, 146.7, 151, 1000],
                    11: [131.8, 136.2, 140.2, 145.3, 148.8, 153.2, 157.7, 1000],
                    12: [137.6, 142.2, 145.9, 150.4, 154.2, 159.2, 163.2, 1000],
                    13: [143, 148.3, 151.8, 155.5, 159.8, 163.7, 168, 1000],
                    14: [147.8, 152.6, 155.4, 159, 163.6, 167.2, 171.2, 1000],
                    15: [150.7, 154.4, 157.2, 161.2, 166, 169.2, 173.4, 1000],
                    16: [151.6, 155.2, 158, 162.5, 166.8, 170.2, 173.8, 1000],
                    17: [152.2, 155.8, 158.6, 162.8, 169.2, 170.4, 174.2, 1000]}
            }
        }

        indicators = {
            '0-3': {
                'br': (22, 28),
                'hr': (80, 100),
                'bp': (90, 100, 60, 70)},
            '3-6': {
                'br': (20, 28),
                'hr': (80, 100),
                'bp': (96, 110, 60, 70)},
            '6-12': {
                'br': (18, 22),
                'hr': (70, 90),
                'bp': (100, 110, 60, 75)},
            '>12': {
                'br': (18, 22),
                'hr': (70, 80),
                'bp': (110, 120, 70, 78)},
        }

        if age <= 3:
            indicator = indicators['0-3']
        elif age <= 6:
            indicator = indicators['3-6']
        elif age <= 12:
            indicator = indicators['6-12']
        else:
            indicator = indicators['>12']

        render_data['temp'] = random.choice(['36,6', '36,7', '36,5'])
        render_data['br'] = random.randrange(start=indicator['br'][0], stop=indicator['br'][1], step=2)
        render_data['hr'] = random.randrange(start=indicator['hr'][0], stop=indicator['hr'][1], step=2)
        render_data['bp'] = f"{random.randrange(start=indicator['bp'][0], stop=indicator['bp'][1], step=1)}/" \
                            f"{random.randrange(start=indicator['bp'][2], stop=indicator['bp'][3], step=1)}"

        anthro = ' (выше- ниже-) среднее, (дис-) гармоничное'
        try:
            height = float(render_data.get('height'))
            weight = float(render_data.get('weight'))
            if render_data.get('gender').lower().startswith('м'):
                anthro_height = anthropometry['мужской']['height'].get(age, [])
                anthro_weight = anthropometry['мужской']['weight'].get(age, [])
            else:
                anthro_height = anthropometry['женский']['height'].get(age, [])
                anthro_weight = anthropometry['женский']['weight'].get(age, [])
            for a_height in anthro_height:
                if height < a_height:
                    index_height = anthro_height.index(a_height)
                    break
            else:
                index_height = 7

            for a_weight in anthro_weight:
                if weight <= a_weight:
                    index_weight = anthro_weight.index(a_weight)
                    break
            else:
                index_weight = 7
            if index_height == 0:
                anthro = 'Низкое '
            elif index_height <= 2:
                anthro = 'Ниже среднего '
            elif index_height <= 4:
                anthro = 'Среднее '
            elif index_height <= 6:
                anthro = 'Выше среднего '
            elif index_height == 7:
                anthro = 'Высокое '

            if abs(index_weight - index_height) <= 1:
                anthro += 'гармоничное'
            elif abs(index_weight - index_height) < 3:
                anthro += 'дисгармоничное'
            else:
                anthro += 'резко дисгармоничное'

            if 'Физическое развитие (выше- ниже-) среднее, (дис-) гармоничное' in render_data.get('diagnosis'):
                render_data['diagnosis'] = \
                    render_data.get('diagnosis').replace('Физическое развитие (выше- ниже-) '
                                                         'среднее, (дис-) гармоничное',
                                                         f"Физическое развитие: {anthro}")
        except Exception:
            pass
        render_data['anthro'] = anthro

        render_data['imt'] = round(float(render_data.get('weight')) /
                                   (float(render_data.get('height')) / 100) ** 2, 1)

        render_data['additional_medical_information'] = render_data.get('additional_medical_information',
                                                                        '').replace(
            'АД ________', f"АД {render_data['bp']} мм.рт.ст.")
        render_data['diagnosis'] = render_data.get('diagnosis', '').replace(
            'школе с _______ лет', f"школе с {age} лет")

    if (type_certificate == "Оформление в ДДУ / СШ / ВУЗ" and "Детское Дошкольное Учреждение"
            not in render_data.get('place_of_requirement')) or type_certificate == 'Об усыновлении (удочерении)':
        doc_name = ""
        if type_certificate == 'Оформление в ДДУ / СШ / ВУЗ':
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} " \
                       f"справка Оформление.docx"
            if not render_data.get('number_cert'):
                render_data['number_cert'] = '№ ______'
        elif data['certificate'].get('type_certificate') == 'Об усыновлении (удочерении)':
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} " \
                       f"справка Об усыновлении.docx"

        if 'Для поступления в учреждения высшего' in render_data.get('place_of_requirement'):
            render_data['name'] += f'\nИдентификационный № _______________________________'
            age = get_age(data['patient'].get('birth_date'))
            if age >= 17:
                render_data['additional_medical_information'] += '\nФлюорография: № _________ от __ . __ . ____ ' \
                                                                 'Заключение: ОГК без патологии'
            else:
                render_data['additional_medical_information'] += '\nФлюорография: не подлежит по возрасту.'

        manager = data["doctor"].get('manager')
        if manager:
            render_data['manager'] = manager
        else:
            render_data['manager'] = '______________________'

        doc = DocxTemplate(f".{os.sep}example{os.sep}certificate{os.sep}справка а4 годовая.docx")
        doc.render(render_data)
        doc.save(doc_name)

        if vaccinations.create_vaccination(user_id=data['patient'].get('amb_cart'), size=4):
            master = Document(doc_name)
            master.add_page_break()
            composer = Composer(master)
            doc_temp = Document(f'.{os.sep}generated{os.sep}прививки.docx')
            composer.append(doc_temp)
            composer.save(doc_name)

        # file = open(doc_name, 'rb')
        # subprocess.Popen(doc_name)
        os.system(f"start {doc_name}")


        doc = DocxTemplate(f".{os.sep}example{os.sep}certificate{os.sep}осмотр.docx")
        doc.render(render_data)
        doc.save(f".{os.sep}generated{os.sep}{data['certificate'].get('name').split()[0]} осмотр.docx")
        # file = open(f".{os.sep}generated{os.sep}{data['certificate'].get('name').split()[0]} осмотр.docx", 'rb')
        os.system(f"start .{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} осмотр.docx")

        subprocess.Popen(f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} осмотр.docx")

    else:

        if type_certificate.startswith('ЦКРОиР'):
            doc = DocxTemplate(f".{os.sep}example{os.sep}certificate{os.sep}выписка ЦКРОиР.docx")
            doc.render(render_data)
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} выписка ЦКРОиР.docx"
            doc.save(doc_name)

        elif type_certificate.startswith('Бесплатное питание'):
            doc = DocxTemplate(f".{os.sep}example{os.sep}certificate{os.sep}Выписка.docx")
            doc.render(render_data)
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} Выписка.docx"
            doc.save(doc_name)

        elif data['certificate'].get('type_certificate') in ('Годовой медосмотр', 'В детский лагерь'):
            if data['certificate'].get('type_certificate').startswith('В детский лагерь'):

                info = (data['doctor'].get('doctor_district'),
                        None,
                        datetime.now().strftime("%d.%m.%Y"),
                        render_data.get('name'),
                        render_data.get('birth_date'),
                        render_data.get('gender'),
                        render_data.get('address')
                        )
                number = data_base.save_certificate_ped_div(data=info,
                                                            type_table='certificate_ped_div',
                                                            district_pd=data['doctor'].get('ped_div'))

                # save_certificate_ped_div(data=info, type_table='certificate_camp')
                render_data['number_cert'] = f"№ {data['certificate'].get('doctor_district')} / {number}"
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} " \
                       f"справка {type_certificate}.docx"

            master = Document(f".{os.sep}example{os.sep}certificate{os.sep}справка а5.docx")
            master.add_page_break()
            composer = Composer(master)

            if type_certificate.startswith('В детский лагерь'):
                if vaccinations.create_vaccination(user_id=data['patient'].get('amb_cart'),
                                                   size=5):
                    doc_temp = Document(f'.{os.sep}generated{os.sep}прививки.docx')
                    composer.append(doc_temp)
                    master.add_page_break()

            master.add_section()
            master.sections[-1].orientation = WD_ORIENT.LANDSCAPE
            master.sections[-1].page_width = master.sections[0].page_height
            master.sections[-1].page_height = master.sections[0].page_width
            doc_temp = Document(f".{os.sep}example{os.sep}certificate{os.sep}осмотр.docx")
            composer.append(doc_temp)

            composer.save(doc_name)

            doc = DocxTemplate(doc_name)
            doc.render(render_data)

            doc.save(doc_name)

        else:
            doc = DocxTemplate(f".{os.sep}example{os.sep}certificate{os.sep}справка а5.docx")
            doc.render(render_data)
            doc_name = f".{os.sep}generated{os.sep}{data['patient'].get('name').split()[0]} " \
                       f"справка {type_certificate}.docx"
            doc.save(doc_name)

            if (data['certificate'].get('type_certificate') in ('В детский лагерь',
                                                                'Может работать по специальности...') or
                    (data['certificate'].get('type_certificate') == 'Об отсутствии контактов' and
                     data['certificate'].get('place_of_requirement') == 'В стационар')):

                if vaccinations.create_vaccination(user_id=data['patient'].get('amb_cart'), size=5):
                    master = Document(doc_name)
                    master.add_page_break()
                    composer = Composer(master)
                    doc_temp = Document(f'.{os.sep}generated{os.sep}прививки.docx')
                    composer.append(doc_temp)
                    composer.save(doc_name)

        os.system(f"start {doc_name}")
        # subprocess.Popen(doc_name)

