from docx import Document
from docx.shared import Cm
from docx.shared import Pt
import os
from tkinter import messagebox


import sqlite3 as sq


def create_doc(amb_cart):
    if create_vaccination(amb_cart, 5):
        os.system(f"start .{os.sep}generated{os.sep}прививки.docx")

    else:
        messagebox.showinfo('Ошибка!', 'Не удалось создать прививки!')


def create_vaccination(user_id, size):
    try:
        if size == 5:
            size = 1
        elif size == 4:
            size = 2

        with sq.connect('patient_data_base.db') as conn:
            cur = conn.cursor()
            cur.execute(f"SELECT Прививки_шапка, Прививки FROM patient_data WHERE amb_cart LIKE '{user_id}'")
            user_data, vaccination = cur.fetchone()

        document = Document()

        paragraph = document.add_paragraph()
        p = paragraph.add_run(f'{user_data}')
        r_fmt = p.font
        r_fmt.name = 'Times New Roman'
        if size == 2:
            r_fmt.size = Pt(12)
        else:
            r_fmt.size = Pt(10)

        info = ["Возраст",
                "Дата проведения прививки",
                "Тип иммуни-\nзации",
                "Наименование препарата",
                "Страна изготови-\nтель препарата",
                "Доза",
                "Серия",
                "Реакция",
                "Мед. отвод"]

        all_tabs = (
            'Прививки против туберкулёза',
            'Прививки против гепатита В',
            'Прививки против кори, эпидемического паротита и краснухи',
            'Прививки против полиомиелита',
            'Прививки против дифтерии, коклюша, столбняка',
            'Прививки против других инфекций')

        table = document.add_table(rows=(len(vaccination.split('\n')) + 1), cols=9)
        table.style = 'Table Grid'
        # if size == 2:
        #     widths = (Cm(1.5*1.5), Cm(1.8*1.5), Cm(1.5*1.5), Cm(1.5*1.5), Cm(1.5*1.5), Cm(1.5*1.5), Cm(1.5*1.5), Cm(1.5*1.5), Cm(1.4*1.5))
        # else:
        #     widths = (Cm(1.5), Cm(1.8), Cm(1.2), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.4))
        # for row in table.rows:
        #     for idx, width in enumerate(widths):
        #         row.cells[idx].width = width

        hdr_cells = table.rows[0].cells
        for i in range(9):
            hdr_cells[i].text = info[i]

            rc = hdr_cells[i].paragraphs[0].runs[0]
            rc.font.name = 'Times New Roman'
            # rc.font.bold = True
            if size == 2:
                rc.font.size = Pt(10)
            else:
                rc.font.size = Pt(7)

        info = vaccination.split('\n')

        for i in range(1, len(info) + 1):
            hdr_cells = table.rows[i].cells
            if info[i - 1] in all_tabs:
                hdr_cells[0].text = info[i - 1]
                rc = hdr_cells[0].paragraphs[0].runs[0]
                if size == 2:
                    rc.font.size = Pt(10)
                else:
                    rc.font.size = Pt(8)

                hdr_cells[0].merge(hdr_cells[8])
            else:
                loc_info = info[i - 1].split('__')
                for q in range(9):
                    hdr_cells[q].text = loc_info[q]

                    rc = hdr_cells[q].paragraphs[0].runs[0]
                    rc.font.name = 'Times New Roman'
                    if size == 2:
                        rc.font.size = Pt(10)
                    else:
                        rc.font.size = Pt(8)

        if size == 2:
            widths = (
                Cm(1.5 * 1.5), Cm(1.8 * 1.5), Cm(1.5 * 1.5), Cm(1.5 * 1.5), Cm(1.5 * 1.5), Cm(1.5 * 1.5), Cm(1.5 * 1.5),
                Cm(1.5 * 1.5), Cm(1.4 * 1.5))
        else:
            widths = (Cm(1.6), Cm(2.0), Cm(1.0), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.4))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            if size == 1:
                section.page_height = Cm(21)
                section.page_width = Cm(14.8)

        document.save(f'.{os.sep}generated{os.sep}прививки.docx')

        if info:
            return True

        else:
            return False

    except Exception as ex:
        print('Exception create_vaccination: ', ex)
        return False
