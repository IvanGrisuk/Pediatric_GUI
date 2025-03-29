import threading

import shutil
import time

import os

from tkinter import *
from tkinter import Label, Frame
import tkinter as tk
from PIL import Image, ImageTk

from datetime import datetime

from docx.enum.section import WD_ORIENT

from database import data_base
# from examination import paste_examination


from variables import all_patient, app_info, user, program_version
from add_func.search_patient import search_patient
from add_func.title_function import redact_doctor, redact_patient


def keypress(event):
    if event.keycode == 86 or event.keycode == 150994966:
        event.widget.event_generate('<<Paste>>')
    elif event.keycode == 67 or event.keycode == 134217731:
        event.widget.event_generate('<<Copy>>')
    elif event.keycode == 88 or event.keycode == 117440536:
        event.widget.event_generate('<<Cut>>')


def main_root():
    def start_action(func=None):
        animation = app_info['title_frame'].get('animation')

        def check_thread(thread_):
            if thread_.is_alive():
                animation.set(animation.get()[-1] + animation.get()[:-1])
                root.after(50, lambda: check_thread(thread))
            else:
                animation.set("")

        def run_action():
            if func:
                func()
            else:
                time.sleep(5)

        animation.set("▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒░")
        thread = threading.Thread(target=run_action)
        thread.start()
        check_thread(thread)

    def finish():
        root.destroy()  # ручное закрытие окна и всего приложения
        print("Закрытие приложения")

    def edit_local_db():
        last_edit_srv = data_base('last_edit_patient_db_srv')
        if last_edit_srv:
            load_info_text.set("Соединение с сервером установлено")
        else:
            load_info_text.set("Ошибка подключения к базе данных!")
        if not last_edit_srv:
            user['error_connection'] = True
        else:
            last_edit_loc = data_base('last_edit_patient_db_loc')
            if last_edit_loc != last_edit_srv:
                load_info_text.set("Обнаружена новая версия базы данных пациентов\n"
                                   "Начинаю обновление...")
                shutil.copy2(f"{user['app_data'].get('path_srv_data_base')}patient_data_base.db",
                             f".{os.sep}data_base{os.sep}patient_data_base.db")
                load_info_text.set(f"База данных пациентов обновлена")
            else:
                load_info_text.set(f"У вас актуальная версия базы данных")
        if not all_patient:
            load_info_text.set(f"Создание базы данных пациентов")
            data_base('select_all_patient')
            load_info_text.set(f"Загрузка завершена")

        load_info_text.set("Синхронизация осмотров...")

        app_info['load_info_text'] = load_info_text
        answer = data_base(command='examination__edit_examination_loc')

        if 'Exception' in answer:
            load_info_text.set(f"{answer}")
            time.sleep(1)

        load_info_text.set("")
        crynet_systems_label.bind('<Double-Button-1>', start_edit_local_db)

    def start_edit_local_db(event=None):
        crynet_systems_label.unbind('<Double-Button-1>')
        start_action(edit_local_db)

    def paste_log_in_root():
        def paste_frame_title():
            app_info['title_frame']['doctor_info_text'] = doctor_info_text
            app_info['title_frame']['animation'] = animation
            app_info['title_frame']['load_info_text'] = load_info_text

            frame_main = Frame(log_in_root, bg="#36566d")

            frame = Frame(frame_main, bg="#36566d")

            Label(frame, textvariable=doctor_info_text,
                  font=('Comic Sans MS', user.get('text_size')),
                  bg="#36566d", fg='white', anchor='nw', justify='left'
                  ).pack(fill='x', expand=True, padx=2, pady=4)
            frame.pack(padx=2, pady=2, side='left', anchor='nw')

            frame = Frame(frame_main, bg="#36566d")

            crynet_systems_label = Label(frame, image=image_crynet_systems,
                                         anchor='ne', bg="#36566d", fg='white', )
            crynet_systems_label.pack(fill='both', expand=True, padx=2, pady=2)
            app_info['title_frame']['crynet_systems_label'] = crynet_systems_label

            Label(frame, textvariable=load_info_text,
                  anchor='ne', bg="#36566d", fg='white', compound='bottom'
                  ).pack(fill='x', expand=True, padx=2, pady=2, side='left')
            Label(frame, textvariable=animation,
                  anchor='ne', bg="#36566d", fg='white', compound='bottom'
                  ).pack(padx=2, pady=2, side='left')

            frame.pack(fill='x', expand=True, padx=2, pady=2, anchor='ne')
            frame_main.pack(fill='x', expand=True, padx=2, pady=2, side='top', anchor='n')

        def select_doctor_name():
            doctor_info_text.set(f"Выбран доктор: "
                                 f"{app_info['all_doctor_info'][int(selected_doctor_name.get())].get('doctor_name')}")

            if app_info['all_doctor_info'][int(selected_doctor_name.get())].get('password'):
                frame_pass.pack_configure(fill='both', expand=True, padx=2, pady=2)
                pass_txt.focus()
            else:
                open_main_root()
                frame_pass.pack_forget()

        def open_main_root():
            frame_pass.destroy()

            if not user.get('error_connection'):
                data_base('activate_app')

            doctor_name = int(selected_doctor_name.get())
            user['doctor_name'] = app_info['all_doctor_info'][doctor_name].get('doctor_name')
            user['doctor_id'] = app_info['all_doctor_info'][doctor_name].get('doctor_id')
            user['password'] = app_info['all_doctor_info'][doctor_name].get('password')
            user['doctor_district'] = app_info['all_doctor_info'][doctor_name].get('district')
            user['ped_div'] = app_info['all_doctor_info'][doctor_name].get('ped_div')
            user['manager'] = app_info['all_doctor_info'][doctor_name].get('manager')
            user['admin_status'] = app_info['all_doctor_info'][doctor_name].get('admin_status')
            user['specialities'] = app_info['all_doctor_info'][doctor_name].get('specialities')
            user['add_info'] = app_info['all_doctor_info'][doctor_name].get('add_info')

            user['my_saved_diagnosis'] = app_info['all_doctor_info'][doctor_name].get('my_saved_diagnosis')
            user['my_LN'] = app_info['all_doctor_info'][doctor_name].get('my_LN')
            user['my_sport_section'] = app_info['all_doctor_info'][doctor_name].get('my_sport_section')

            paste_frame_main()

        def connect_to_srv_data_base():
            load_info_text.set(f"Попытка подключения к базе данных...")

            status = data_base('get_all_doctor_info')
            if status == 'srv':
                load_info_text.set("Соединение с сервером установлено")
                data_base('activate_app')
            else:
                load_info_text.set("Ошибка подключения к базе данных!")

            if user['app_data'].get('last_reg_password'):
                if (datetime.now() - datetime.strptime(user['app_data'].get('last_reg_password'),
                                                       "%Y-%m-%d %H:%M:%S.%f")).days > 60:
                    load_info_text.set('Срок активации истек! '
                                       '\nВведите пароль для продления 60-дневной подписки')

                    frame_pass.pack_configure(fill='both', expand=True, padx=2, pady=2)
                    app_info['check_pass_app'] = True

            if not app_info.get('check_pass_app'):
                if app_info.get('all_doctor_info'):
                    frame_doc = Frame(log_in_root, borderwidth=1, relief="solid", padx=8, pady=10)

                    doctor_info_text = app_info['title_frame'].get('doctor_info_text')
                    doctor_info_text.set("")

                    users_sorted_pd = dict()

                    for doctor_id in sorted(app_info.get('all_doctor_info'),
                                            key=lambda i: app_info['all_doctor_info'][i].get('doctor_name')):
                        ped_div = app_info['all_doctor_info'][doctor_id].get('ped_div')
                        doctor_name = app_info['all_doctor_info'][doctor_id].get('doctor_name')
                        if ped_div not in users_sorted_pd:
                            users_sorted_pd[ped_div] = [[]]
                        if len(users_sorted_pd.get(ped_div)[-1]) == 5:
                            users_sorted_pd[ped_div].append([])
                        users_sorted_pd[ped_div][-1].append((doctor_name, doctor_id))

                    Label(frame_doc, text="Выберите учетную запись",
                          bg="#36566d", fg='white',
                          font=('Comic Sans MS', user.get('text_size'))
                          ).pack(fill='x', expand=True, padx=2, pady=2, ipady=2)

                    for ped_div in sorted(users_sorted_pd):
                        if ped_div.isdigit():
                            text = f'{ped_div}-е ПО'
                        else:
                            text = f'{ped_div}'
                        Label(frame_doc, text=text,
                              font=('Comic Sans MS', user.get('text_size')),
                              bg='white').pack(fill='x', expand=True, padx=2, pady=2, anchor='n')
                        for list_doc in users_sorted_pd.get(ped_div):
                            frame = Frame(frame_doc)
                            for doctor_name, doctor_id in list_doc:
                                Radiobutton(master=frame, text=doctor_name,
                                            font=('Comic Sans MS', user.get('text_size')),
                                            command=select_doctor_name,
                                            value=doctor_id, variable=selected_doctor_name,
                                            indicatoron=False, selectcolor='#77f1ff'
                                            ).pack(fill='both', expand=True, padx=2, pady=2, side='left')
                            frame.pack(fill='x', expand=True, padx=2, pady=2, anchor='n')

                    frame_doc.pack(fill='both', expand=True, padx=2, pady=2)

        def is_valid__password(password=None):
            password = txt_password_variable.get()
            if app_info.get('check_pass_app'):
                if password == "profkiller_10539008":
                    text_is_correct_password.set('Пароль принят')
                    app_info['check_pass_app'] = False
                    data_base('activate_app')
                    log_in_root.destroy()
                    paste_log_in_root()
                else:
                    text_is_correct_password.set('Пароль не верен!')
            else:
                if password == app_info['all_doctor_info'][int(selected_doctor_name.get())].get('password'):
                    text_is_correct_password.set('Пароль принят')
                    open_main_root()
                else:
                    text_is_correct_password.set('Пароль не верен!')

            return True

        if app_info.get('frame_main'):
            app_info['frame_main'].destroy()
            app_info['frame_main'] = None


        selected_doctor_name = StringVar()
        txt_password_variable = StringVar()
        text_is_correct_password = StringVar()

        doctor_info_text = StringVar()
        animation = StringVar()
        load_info_text = StringVar()

        doctor_info_text.set("")
        animation.set("")
        load_info_text.set("")

        log_in_root = Frame(master=root, bg="#36566d")
        log_in_root.pack(fill='both', expand=True, padx=2, pady=2)
        app_info['log_in_root'] = log_in_root
        user['error_connection'] = False

        load_info_text.set('Запуск программы...')

        frame_pass = Frame(log_in_root, borderwidth=1, relief="solid", padx=8, pady=10)

        Label(frame_pass, text='Введите пароль: ',
              font=('Comic Sans MS', user.get('text_size')), bg='white'
              ).pack(fill='both', expand=True, padx=2, pady=2, side='left')
        pass_txt = Entry(frame_pass, width=40,
                         font=('Comic Sans MS', user.get('text_size')),
                         justify="center",
                         textvariable=txt_password_variable,
                         show="*"
                         )
        pass_txt.pack(fill='both', expand=True, padx=2, pady=2, side='left')
        pass_txt.bind('<Return>', is_valid__password)

        Button(frame_pass, text='Подтвердить', command=is_valid__password,
               font=('Comic Sans MS', user.get('text_size'))
               ).pack(fill='both', expand=True, padx=2, pady=2, side='left')

        Label(frame_pass, textvariable=text_is_correct_password,
              font=('Comic Sans MS', user.get('text_size')), bg='white', foreground="red"
              ).pack(fill='both', expand=True, padx=2, pady=2, side='left')

        log_in_root.columnconfigure(index='all', minsize=40, weight=1)
        log_in_root.rowconfigure(index='all', minsize=20)

        root.geometry(f"+50+50")
        paste_frame_title()
        start_action(connect_to_srv_data_base)

    def paste_frame_main():

        def paste_frame_title():
            def pack_frame_butt():
                def select_type_cert():
                    pass

                frame_main_but = Frame(frame_title)
                all_butt = ('Справка', 'Анализы', 'Осмотры',
                            'Осмотры до года', 'Вкладыши', 'Прививки', 'Направления', 'Мой прием')
                for butt in all_butt:
                    Radiobutton(master=frame_main_but,
                                text=butt,
                                font=('Comic Sans MS', user.get('text_size')),
                                value=butt,
                                variable=selected_button,
                                command=select_type_cert,
                                indicatoron=False, bg='#f0fffe', selectcolor='#77f1ff'
                                ).pack(fill='both', expand=True, padx=1, pady=1, side='left')
                frame_main_but.pack(fill='both', expand=True, padx=2, pady=2)

            def pack_menu():

                root.option_add("*tearOff", FALSE)

                main_menu = Menu()

                file_menu = Menu()
                file_menu.add_command(label="Редактировать мои данные", command=redact_doctor)
                file_menu.add_command(label="Редактировать пациента", command=redact_patient)
                file_menu.add_command(label="Добавить нового пациента", command=add_new_patient)
                file_menu.add_separator()
                file_menu.add_command(label="Выйти", command=paste_log_in_root)

                main_menu.add_cascade(label="Меню", menu=file_menu)
                # main_menu.add_cascade(label="Edit")
                # main_menu.add_cascade(label="View")

                root.config(menu=main_menu)

            def pack_lbl():
                frame_lbl = Frame(frame_title, bg="#36566d")
                frame_1 = Frame(frame_lbl, bg="#36566d")
                Label(frame_1, textvariable=doctor_info_text,
                      font=('Comic Sans MS', user.get('text_size')),
                      bg="#36566d", fg='white', anchor='nw', justify='left'
                      ).pack(fill='x', expand=True, padx=2, pady=2)
                frame = Frame(frame_1, bg="#36566d")

                Label(frame, textvariable=patient_info_text,
                      font=('Comic Sans MS', user.get('text_size')),
                      bg="#36566d", fg='white', anchor='nw', justify='left',
                      ).pack(fill='both', expand=True, padx=2, side='left')

                Button(frame, text='Поиск пациента', command=search_patient,
                       font=('Comic Sans MS', user.get('text_size'))
                       ).pack(fill='x', anchor='w', padx=2, side='left')


                frame.pack(padx=2, pady=2, anchor='nw')
                frame_1.pack(padx=2, pady=2, anchor='nw', side='left')

                frame = Frame(frame_lbl, bg="#36566d")

                crynet_systems_label = Label(frame, image=image_crynet_systems,
                                             anchor='ne', bg="#36566d", fg='white', )
                crynet_systems_label.pack(fill='both', expand=True, padx=2, pady=2)
                app_info['title_frame']['crynet_systems_label'] = crynet_systems_label

                Label(frame, textvariable=load_info_text,
                      anchor='ne', bg="#36566d", fg='white', compound='bottom'
                      ).pack(fill='x', expand=True, padx=2, pady=2, side='left')
                Label(frame, textvariable=animation,
                      anchor='ne', bg="#36566d", fg='white', compound='bottom'
                      ).pack(padx=2, pady=2, side='left')

                frame.pack(fill='x', expand=True, padx=2, pady=2, anchor='ne')
                frame_lbl.pack(fill='x', expand=True)

            frame_title = Frame(frame_main, bg="#36566d")

            doctor_info_text = StringVar()
            patient_info_text = StringVar()
            animation = StringVar()
            load_info_text = StringVar()

            patient_info_text.set("Пациент:")
            animation.set("")
            load_info_text.set("")
            doctor_info_text.set(f"Учетная запись: "
                                 f"Доктор: {user.get('doctor_name').upper()}    "
                                 f"Зав: {user.get('manager')};    "
                                 f"Участок: {user.get('doctor_district')};    "
                                 f"ПО: {user.get('ped_div')}")

            app_info['title_frame']['doctor_info_text'] = doctor_info_text
            app_info['title_frame']['patient_info_text'] = patient_info_text
            app_info['title_frame']['animation'] = animation
            app_info['title_frame']['load_info_text'] = load_info_text
            app_info['title_frame']['frame_title'] = frame_title

            frame_title.pack(fill='x', expand=True, padx=2, pady=2, side='top', anchor='n')

            pack_menu()
            pack_lbl()
            pack_frame_butt()


        # def pack_main_frame():
        #
        #     app_info['frame_main'] = frame_main
        #
        #     crynet_systems_label.bind('<Double-Button-1>', start_edit_local_db)
        #
        #
        #     # frame_buttons = Frame(master=frame_main, bg="#36566d", padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Справка', command=fast_certificate,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Анализы', command=analyzes_cmd,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Вкладыши', command=blanks_cmd,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Прививки', command=vaccination_cmd,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Направления', command=direction_cmd,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Мой прием', command=open_last_examination,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Осмотры', command=examination_cmd,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Осмотры до года', command=examination_cmd_child,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # if 'local_admin' in str(user.get('add_info', "")):
        #     #     Button(frame_buttons, text='Журнал справок', command=download_ped_div,
        #     #            font=('Comic Sans MS', user.get('text_size'))
        #     #            ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # Button(frame_buttons, text='Сменить пользователя', command=change_account,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     # Button(frame_buttons, text='Редактировать данные', command=redact_doctor,
        #     #        font=('Comic Sans MS', user.get('text_size'))
        #     #        ).pack(fill='both', expand=True, padx=2, pady=2)
        #     #
        #     # frame_buttons.pack(fill='both', expand=True, padx=2, pady=2)
        #
        #     # frame_main.pack(fill='both', padx=2, pady=2, anchor='n')
        #     frame_main.pack(fill='both', padx=2, pady=2, side='top', anchor='n')

        # def add_new_patient():
        #     def save():
        #         def check_input():
        #             error_flag = False
        #             for marker in local_data:
        #                 if marker in ("№ участка", "Фамилия", "Имя", "Отчество", "Пол", "Дата рождения",
        #                               "Адрес") and not local_data.get(marker).get():
        #                     messagebox.showerror('Ошибка', f"Ошибка!\nНе указан пункт\n'{marker}'")
        #                     break
        #                 elif marker in ("№ амбулаторной карты", "№ участка") and local_data.get(
        #                         marker).get() and not local_data.get(marker).get().isdigit():
        #                     messagebox.showerror('Ошибка', f"Ошибка!\nУкажите пункт\n'{marker}'\nчислом")
        #                     break
        #                 elif marker == "Дата рождения":
        #                     try:
        #                         if (datetime.now() - datetime.strptime(local_data.get(marker).get(),
        #                                                                "%d.%m.%Y")).days < 0:
        #                             messagebox.showerror('Ошибка', f"Дата рождения не может быть больше текущей даты!")
        #                             break
        #                     except Exception:
        #                         messagebox.showerror('Ошибка', f"Дата рождения должна быть в формате 'ДД.ММ.ГГГГ'")
        #                         break
        #             else:
        #                 return True
        #             return False
        #
        #         if check_input():
        #             insert_data = list()
        #             for marker in ("№ участка", "№ амбулаторной карты",
        #                            "Фамилия", "Имя", "Отчество", "Пол",
        #                            "Дата рождения", "Адрес",
        #                            "None", "None", "None"):
        #                 if local_data.get(marker) and local_data.get(marker).get().strip():
        #                     insert_data.append(local_data.get(marker).get().strip())
        #                 else:
        #                     insert_data.append("")
        #             if data_base(command='save_new_patient', insert_data=insert_data):
        #                 messagebox.showinfo('Инфо', "Данные успешно сохранены!")
        #                 patient['name'] = f"{local_data.get('Фамилия').get().strip()} " \
        #                                   f"{local_data.get('Имя').get().strip()} " \
        #                                   f"{local_data.get('Отчество').get().strip()}".strip()
        #                 patient['birth_date'] = f"{local_data.get('Дата рождения').get().strip()}"
        #                 patient['gender'] = f"{local_data.get('Пол').get().strip()}"
        #                 patient['amb_cart'] = f"{local_data.get('№ амбулаторной карты').get().strip()}"
        #                 patient['district'] = f"{local_data.get('№ участка').get().strip()}"
        #                 patient['address'] = f"{local_data.get('Адрес').get().strip()}"
        #                 patient['age'] = get_age_d_m_y(patient.get('birth_date'))
        #
        #                 patient_info.set(f"ФИО: {patient.get('name')}\t"
        #                                  f"Дата рождения: {patient.get('birth_date')}    {patient['age'].get('age_txt')}\n"
        #                                  f"Адрес: {patient.get('address')}\n"
        #                                  f"№ амб: {patient.get('amb_cart')}\t"
        #                                  f"Участок: {patient.get('district')}")
        #
        #                 new_root.destroy()
        #             else:
        #                 messagebox.showerror('Ошибка', f"Ошибка!\nОшибка сохранения двнных")
        #
        #     new_root = Toplevel()
        #     new_root.title('Добавление нового пациента')
        #     root.bind("<Control-KeyPress>", keypress)
        #     local_data = {
        #         "№ амбулаторной карты": StringVar(),
        #         "№ участка": StringVar(),
        #         "Фамилия": StringVar(),
        #         "Имя": StringVar(),
        #         "Отчество": StringVar(),
        #         "Пол": StringVar(),
        #         "Дата рождения": StringVar(),
        #         "Адрес": StringVar(),
        #     }
        #
        #     row = 0
        #     for marker in local_data:
        #
        #         Label(new_root, text=marker,
        #               font=('Comic Sans MS', user.get('text_size')),
        #               bg="#36566d", fg='white').grid(column=0, row=row, sticky='nwse', padx=2, pady=2)
        #         if marker == 'Пол':
        #             combo_sex = Combobox(new_root, font=('Comic Sans MS', user.get('text_size')), state="readonly",
        #                                  textvariable=local_data.get(marker))
        #             combo_sex['values'] = ["", "мужской", "женский"]
        #             combo_sex.current(0)
        #             combo_sex.grid(column=1, row=row, sticky='nwse')
        #         else:
        #
        #             Entry(new_root, width=30, font=('Comic Sans MS', user.get('text_size')),
        #                   textvariable=local_data.get(marker)
        #                   ).grid(column=1, row=row, sticky='nwse', ipadx=2, ipady=2)
        #         row += 1
        #
        #     Button(new_root, text='Сохранить', command=save, font=('Comic Sans MS', user.get('text_size'))).grid(
        #         columnspan=2, sticky='ew')
        #
        #     new_root.mainloop()
        #
        # def download_ped_div():
        #     def start_search():
        #         pediatric_division = user.get('ped_div')
        #         download_ped_div_variable.set(f"Выгрузка справок {pediatric_division} ПО...\n"
        #                                       f"Обращение к базе данных...")
        #         download_ped_div_root.update()
        #
        #         info = data_base(f"get_certificate_for_district__certificate_ped_div__{pediatric_division}")
        #         if not info:
        #             download_ped_div_variable.set(
        #                 f"{download_ped_div_variable.get()} \nОшибка подключения к базе данных")
        #             download_ped_div_root.update()
        #         else:
        #             download_ped_div_variable.set(f"{download_ped_div_variable.get()} ответ получен!\n"
        #                                           f"Создаю документ...")
        #             download_ped_div_root.update()
        #
        #             if pediatric_division == '1':
        #                 document = Document()
        #                 table = document.add_table(rows=(len(info) + 1), cols=10)
        #                 table.style = 'Table Grid'
        #                 # widths = (Cm(1.0), Cm(1.0), Cm(2.0), Cm(3.0), Cm(2.5), Cm(1.8), Cm(3.0))
        #                 # for row in table.rows:
        #                 #     for idx, width in enumerate(widths):
        #                 #         row.cells[idx].width = width
        #                 data_table = ('№ п/п',
        #                               'Дата',
        #                               'ФИО лица, обратившегося за выдачей справки и (или) другого документа',
        #                               'Адрес',
        #                               'Документ, удостоверяющий личность',
        #                               'Наименование справки и (или) другого запрашиваемого документа',
        #                               'Срок исполнения',
        #                               'Размер платы, взимаемый',
        #                               'Дата выдачи справки и (или) другого запрашиваемого документа',
        #                               'ФИО врача')
        #                 hdr_cells = table.rows[0].cells
        #                 for i in range(10):
        #                     hdr_cells[i].text = data_table[i]
        #
        #                     rc = hdr_cells[i].paragraphs[0].runs[0]
        #                     rc.font.name = 'Times New Roman'
        #                     rc.font.size = Pt(10)
        #                     rc.font.bold = True
        #
        #                 len_doc = len(info)
        #                 per_num_data = dict()
        #                 for per_num in (0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1):
        #                     per_num_data[round(len_doc * per_num)] = per_num
        #                 download_ped_div_variable.set(f"{download_ped_div_variable.get()} \n"
        #                                               f"Таблица создана. Заполняю ячейки данными...")
        #                 download_ped_div_root.update()
        #
        #                 for i in range(1, len(info) + 1):
        #                     if per_num_data.get(i):
        #                         download_ped_div_variable.set(f"{download_ped_div_variable.get()} \n"
        #                                                       f"Завершено на {round(per_num_data.get(i) * 100)}%")
        #                         download_ped_div_root.update()
        #
        #                     hdr_cells = table.rows[i].cells
        #                     ped_div, district, num, date, name, birth_date, address, type_cert, doctor_name = info[
        #                         i - 1]
        #                     local_info = (
        #                         num, date, name, address, 'паспорт', type_cert, '1 день', 'бесплатно', date,
        #                         doctor_name)
        #                     for q in range(10):
        #                         hdr_cells[q].text = local_info[q]
        #                         rc = hdr_cells[q].paragraphs[0].runs[0]
        #                         rc.font.name = 'Times New Roman'
        #                         rc.font.size = Pt(9)
        #
        #             else:
        #                 document = Document()
        #                 table = document.add_table(rows=(len(info) + 1), cols=11)
        #                 table.style = 'Table Grid'
        #                 # widths = (Cm(1.0), Cm(1.0), Cm(2.0), Cm(3.0), Cm(2.5), Cm(1.8), Cm(3.0))
        #                 # for row in table.rows:
        #                 #     for idx, width in enumerate(widths):
        #                 #         row.cells[idx].width = width
        #                 data_table = ('№ п/п',
        #                               'ФИО, обратившегося за выдачей справки и (или) другого документа',
        #                               "Дата рождения",
        #                               'Домашний адрес',
        #                               'Дата подачи заявления',
        #                               'Наименование справки и (или) другого запрашиваемого документа',
        #                               'Срок исполнения',
        #                               'Документ, удостоверяющий личность',
        #                               'Размер платы, взимаемой за подачу справки и (или) другого документа',
        #                               'Дата выдачи справки и (или) другого запрашиваемого документа',
        #                               'ФИО врача (роспись заявителя)')
        #                 hdr_cells = table.rows[0].cells
        #                 for i in range(11):
        #                     hdr_cells[i].text = data_table[i]
        #
        #                     rc = hdr_cells[i].paragraphs[0].runs[0]
        #                     rc.font.name = 'Times New Roman'
        #                     rc.font.size = Pt(10)
        #                     rc.font.bold = True
        #                 len_doc = len(info)
        #                 per_num_data = dict()
        #                 for per_num in (0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1):
        #                     per_num_data[round(len_doc * per_num)] = per_num
        #                 download_ped_div_variable.set(f"{download_ped_div_variable.get()} \n"
        #                                               f"Таблица создана. Заполняю ячейки данными...")
        #                 download_ped_div_root.update()
        #
        #                 for i in range(1, len(info) + 1):
        #                     if per_num_data.get(i):
        #                         download_ped_div_variable.set(f"{download_ped_div_variable.get()} \n"
        #                                                       f"Завершено на {round(per_num_data.get(i) * 100)}%")
        #                         download_ped_div_root.update()
        #
        #                     hdr_cells = table.rows[i].cells
        #                     ped_div, district, num, date, name, birth_date, address, type_cert, doctor_name = info[
        #                         i - 1]
        #                     type_cert = f"пункт {type_cert}"
        #                     local_info = (
        #                         num,
        #                         name,
        #                         birth_date,
        #                         address,
        #                         date,
        #                         type_cert,
        #                         '1 день',
        #                         'паспорт',
        #                         'бесплатно',
        #                         date,
        #                         doctor_name)
        #                     for q in range(11):
        #                         hdr_cells[q].text = local_info[q]
        #                         rc = hdr_cells[q].paragraphs[0].runs[0]
        #                         rc.font.name = 'Times New Roman'
        #                         rc.font.size = Pt(9)
        #
        #             sections = document.sections
        #             for section in sections:
        #                 section.orientation = WD_ORIENT.LANDSCAPE
        #                 section.top_margin = Cm(1.5)
        #                 section.bottom_margin = Cm(1.5)
        #                 section.left_margin = Cm(1.5)
        #                 section.right_margin = Cm(1.5)
        #                 section.page_height = Cm(21)
        #                 section.page_width = Cm(29.7)
        #
        #             file_name = f'.{os.sep}generated{os.sep}БРЕД_{pediatric_division}_го_ПО.docx'
        #             file_name = save_document(doc=document, doc_name=file_name)
        #             os.system(f"start {file_name}")
        #             download_ped_div_root.destroy()
        #
        #     download_ped_div_root = Toplevel()
        #     download_ped_div_root.title('Инфо')
        #     download_ped_div_variable = StringVar()
        #     Label(download_ped_div_root, textvariable=download_ped_div_variable,
        #           font=('Comic Sans MS', user.get('text_size')),
        #           bg="#36566d", fg='white').pack(fill='x', expand=True, ipady=3, ipadx=3)
        #     download_ped_div_variable.set("Выгрузка справок в среднем занимает около 5 минут\n"
        #                                   "На время формирования документа не закрывайте приложение\n"
        #                                   "Для начала выгрузки нажмите кнопку 'Начать поиск'")
        #     btn_start = Button(download_ped_div_root, text='Начать поиск', command=start_search,
        #                        font=('Comic Sans MS', user.get('text_size')))
        #     btn_start.pack(fill='x', expand=True, ipady=3, ipadx=3)
        #
        #     download_ped_div_root.mainloop()
        #
        # def download_camp():
        #     def start_search():
        #         district = user.get('doctor_district')
        #         download_camp_variable.set(f"Выгрузка справок детского лагеря {district} участка...\n"
        #                                    f"Обращение к базе данных...")
        #         download_camp_root.update()
        #         info = data_base(f"get_certificate_for_district__certificate_camp__{district}")
        #
        #         if not info:
        #             download_camp_variable.set(f"{download_camp_variable.get()} \nОшибка подключения к базе данных")
        #             download_camp_root.update()
        #
        #         else:
        #             download_camp_variable.set(f"{download_camp_variable.get()} ответ получен!\n"
        #                                        f"Создаю документ...")
        #             download_camp_root.update()
        #
        #             document = Document()
        #             table = document.add_table(rows=(len(info) + 1), cols=7)
        #             table.style = 'Table Grid'
        #             widths = (Cm(1.0), Cm(1.0), Cm(2.0), Cm(3.0), Cm(2.5), Cm(1.8), Cm(3.0))
        #             for row in table.rows:
        #                 for idx, width in enumerate(widths):
        #                     row.cells[idx].width = width
        #             data_table = ('Участок', '№ п/п', 'Дата выписки', 'ФИО', 'Дата рождения', 'Пол', 'Адрес')
        #             hdr_cells = table.rows[0].cells
        #             for i in range(7):
        #                 hdr_cells[i].text = data_table[i]
        #
        #                 rc = hdr_cells[i].paragraphs[0].runs[0]
        #                 rc.font.name = 'Times New Roman'
        #                 rc.font.size = Pt(8)
        #                 rc.font.bold = True
        #             len_doc = len(info)
        #             per_num_data = dict()
        #             for per_num in (0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1):
        #                 per_num_data[round(len_doc * per_num)] = per_num
        #             download_camp_variable.set(f"{download_camp_variable.get()} \n"
        #                                        f"Таблица создана. Заполняю ячейки данными...")
        #             download_camp_root.update()
        #
        #             for i in range(1, len(info) + 1):
        #                 if per_num_data.get(i):
        #                     download_camp_variable.set(f"{download_camp_variable.get()} \n"
        #                                                f"Завершено на {round(per_num_data.get(i) * 100)}%")
        #                     download_camp_root.update()
        #
        #                 hdr_cells = table.rows[i].cells
        #                 for q in range(7):
        #                     if info[i - 1][q]:
        #                         hdr_cells[q].text = info[i - 1][q]
        #
        #                         rc = hdr_cells[q].paragraphs[0].runs[0]
        #                         rc.font.name = 'Times New Roman'
        #                         rc.font.size = Pt(8)
        #
        #             sections = document.sections
        #             for section in sections:
        #                 section.top_margin = Cm(1.5)
        #                 section.bottom_margin = Cm(1.5)
        #                 section.left_margin = Cm(1.5)
        #                 section.right_margin = Cm(1.5)
        #                 section.page_height = Cm(21)
        #                 section.page_width = Cm(14.8)
        #
        #             file_name = f'.{os.sep}generated{os.sep}концлагерь_{district}_участка.docx'
        #             file_name = save_document(doc=document, doc_name=file_name)
        #             os.system(f"start {file_name}")
        #             download_camp_root.destroy()
        #
        #     download_camp_root = Toplevel()
        #     download_camp_root.title('Инфо')
        #     download_camp_variable = StringVar()
        #     Label(download_camp_root, textvariable=download_camp_variable,
        #           font=('Comic Sans MS', user.get('text_size')),
        #           bg="#36566d", fg='white').pack(fill='x', expand=True, ipady=3, ipadx=3)
        #     download_camp_variable.set("Выгрузка лагеря в среднем занимает около 5 минут\n"
        #                                "На время формирования документа не закрывайте приложение\n"
        #                                "Для начала выгрузки нажмите кнопку 'Начать поиск'")
        #     btn_start = Button(download_camp_root, text='Начать поиск', command=start_search,
        #                        font=('Comic Sans MS', user.get('text_size')))
        #     btn_start.pack(fill='x', expand=True, ipady=3, ipadx=3)
        #
        #     download_camp_root.mainloop()
        #
        #
        # def redact_doctor():
        #     change_doctor(command='redact')
        #
        # def add_new_doctor():
        #     change_doctor(command='new')
        #
        #
        # def paste_txt_patient_data(event):
        #     if event.keycode == 86 or event.keycode == 150994966:
        #         txt_patient_data.delete(0, 'end')
        #         event.widget.event_generate('<<Paste>>')
        #         search_patient()
        #
        #     elif event.keycode == 67 or event.keycode == 134217731:
        #         event.widget.event_generate('<<Copy>>')
        #     elif event.keycode == 88 or event.keycode == 117440536:
        #         event.widget.event_generate('<<Cut>>')
        #
        #
        #
        # def write_lbl_doc():
        #
        #     lbl_doc_text.set(f"Учетная запись:\n"
        #                      f"Доктор: {user.get('doctor_name')}\n"
        #                      f"Зав: {user.get('manager')};    "
        #                      f"Участок: {user.get('doctor_district')};    "
        #                      f"ПО: {user.get('ped_div')}")
        #
        #     if 'константинова' in user.get('doctor_name'):
        #         lbl_doc_text.set(f"Учетная запись:\n"
        #                          f"Доктор: Яночка Константиновна\n"
        #                          f"Зав: {user.get('manager')};    "
        #                          f"Участок: {user.get('doctor_district')};    "
        #                          f"ПО: {user.get('ped_div')}")
        #
        #     root.update()
        #
        # def delete_txt_patient_data():
        #     txt_patient_data_variable.set('')
        #
        # def update_font_main():
        #     if app_info.get('frame_main'):
        #         app_info['frame_main'].destroy()
        #         app_info['frame_main'] = None
        #
        #     paste_frame_main()
        #
        # def change_account():
        #     paste_log_in_root()


        def pack_scrolled_frame():
            def resize(event=None):
                region = canvas.bbox(tk.ALL)
                canvas.configure(scrollregion=region)

            def on_binds(event):
                canvas.idbind = canvas.bind_all("<MouseWheel>", on_mousewheel)

            def off_binds(event=None):
                canvas.unbind_all("<MouseWheel>")

            def on_mousewheel(event):
                region = canvas.bbox(tk.ALL)
                canvas.configure(scrollregion=region)
                if os.name == 'posix':
                    canvas.yview_scroll(int(-1 * event.delta), "units")
                else:
                    canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

            root.update_idletasks()
            frame_title = app_info['title_frame'].get('frame_title')
            height = (root.winfo_height() - frame_title.winfo_height())
            width = root.winfo_width()
            if root.winfo_screenwidth() < width:
                width = root.winfo_screenwidth()

            master_frame = Frame(frame_main)
            master_frame.pack(fill='both', expand=True, padx=2, pady=2, ipadx=2, ipady=2)

            scroll_x = tk.Scrollbar(master_frame, orient=tk.HORIZONTAL)
            scroll_y = tk.Scrollbar(master_frame, orient=tk.VERTICAL, width=user.get('text_size', 10) * 2)

            canvas = tk.Canvas(master_frame,
                               xscrollcommand=scroll_x.set,
                               yscrollcommand=scroll_y.set, height=height, width=width)
            scroll_x.config(command=canvas.xview)
            scroll_y.config(command=canvas.yview)

            canvas_frame = Frame(canvas)
            app_info['scrolled_frame']['scrolled_frame'] = canvas_frame
            app_info['scrolled_frame']['canvas'] = canvas
            app_info['scrolled_frame']['scroll_x'] = scroll_x

            # canvas['width'] = int(canvas.winfo_geometry().split('x')[0])
            # canvas_frame['width'] = int(canvas.winfo_geometry().split('x')[0])
            canvas.grid(row=0, column=0, sticky="nsew")
            scroll_x.grid(row=1, column=0, sticky="we")
            scroll_y.grid(row=0, column=1, sticky="ns")

            master_frame.rowconfigure(0, weight=1)
            master_frame.columnconfigure(0, weight=1)

            master_frame.bind("<Configure>", resize)
            master_frame.update_idletasks()
            canvas_frame['height'] = height
            canvas_frame['height'] = canvas.winfo_width()

            canvas.bind("<Enter>", on_binds)
            canvas.bind("<Leave>", off_binds)
            root.update_idletasks()

            # canvas.create_window((0, 0), window=canvas_frame, anchor="nw",
            #                      width=canvas.winfo_width())

        def change_account():
            pass

        def add_new_patient():
            pass

        root.geometry(f"{root.winfo_screenwidth() - 50}x{root.winfo_screenheight() - 100}+0+0")
        root.resizable(False, False)

        if app_info.get('log_in_root'):
            app_info['log_in_root'].pack_forget()
            app_info['log_in_root'].destroy()
            app_info['log_in_root'] = None

        # lbl_doc_text = app_info['title_frame'].get('doctor_info_text')
        # patient_info = app_info['title_frame'].get('patient_info_text')
        # crynet_systems_label = app_info['title_frame'].get('crynet_systems_label')

        frame_main = Frame(master=root, bg="#36566d")
        app_info['frame_main'] = frame_main
        frame_main.pack(fill='both', expand=True)

        txt_patient_data_variable = StringVar()
        selected_button = StringVar()

        paste_frame_title()
        pack_scrolled_frame()

        # pack_frame_butt()

    data_base('create_db')
    root = Tk()
    root.title(f"Генератор v_{program_version}")
    root.config(bg="#36566d")
    root.iconbitmap(default=f"img{os.sep}Crynet_systems.ico")
    root.protocol("WM_DELETE_WINDOW", finish)
    root.bind("<Control-KeyPress>", keypress)

    image_crynet_systems = ImageTk.PhotoImage(Image.open('img/Crynet_systems.png').resize((200, 50)))
    calendar_img = ImageTk.PhotoImage(
        Image.open('img/calendar_img.png').resize((user.get('text_size') * 2, user.get('text_size') * 2)))
    user['calendar_img'] = calendar_img

    paste_log_in_root()
    app_info['frame_main_func'] = paste_frame_main
    root.mainloop()


if __name__ == "__main__":
    main_root()
